from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_startup_report():
    doc = Document()

    # --- Cấu hình Style chung ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- TIÊU ĐỀ ---
    title = doc.add_heading('BÁO CÁO: THIẾT KẾ CƠ CẤU TỔ CHỨC DOANH NGHIỆP KHỞI NGHIỆP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thông tin chung
    p = doc.add_paragraph()
    p.add_run('Đề tài: ').bold = True
    p.add_run('Công ty Giải pháp Dữ liệu thông minh (DataFlow AI)\n')
    p.add_run('Quy mô: ').bold = True
    p.add_run('25 nhân sự')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- PHẦN 1: SƠ ĐỒ TỔ CHỨC ---
    doc.add_heading('1. Sơ đồ tổ chức (Organizational Chart)', level=1)
    doc.add_paragraph(
        "Cấu trúc được thiết kế theo dạng Chức năng (Functional Structure), tập trung vào luồng xử lý "
        "từ dữ liệu thô đến sản phẩm AI hoàn chỉnh. Sơ đồ bao gồm Ban giám đốc và 5 phòng ban chuyên biệt."
    )
    
    # Ghi chú về sơ đồ (Vì docx không render được Mermaid trực tiếp, ta để text mô tả)
    doc.add_paragraph("[Chèn hình ảnh Sơ đồ tổ chức tại đây]", style='Intense Quote')

    # --- PHẦN 2: BẢNG PHÂN BỔ NHÂN SỰ ---
    doc.add_heading('2. Bảng phân bổ nhân sự chi tiết', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header bảng
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Bộ phận'
    hdr_cells[1].text = 'Số lượng'
    hdr_cells[2].text = 'Nhiệm vụ chính'
    
    # Dữ liệu bảng
    data = [
        ['Ban Giám đốc (CEO)', '01', 'Quyết định chiến lược, đối ngoại và gọi vốn.'],
        ['Data Science & AI (R&D)', '07', 'Xây dựng thuật toán, Model Training, đánh giá độ chính xác.'],
        ['Kỹ thuật Dữ liệu (Eng)', '06', 'Xây dựng Pipeline, quản lý Database, Cloud Ops.'],
        ['Kinh doanh & Marketing', '05', 'Tìm kiếm đối tác B2B, quảng bá dịch vụ.'],
        ['Chăm sóc Khách hàng', '04', 'Hỗ trợ triển khai giải pháp, giải đáp thắc mắc Dashboard.'],
        ['Tài chính - Hành chính', '02', 'Quản lý ngân sách, lương thưởng và nhân sự.']
    ]
    
    for dept, qty, task in data:
        row_cells = table.add_row().cells
        row_cells[0].text = dept
        row_cells[1].text = qty
        row_cells[2].text = task

    # --- PHẦN 3: CHỨC NĂNG NHIỆM VỤ ---
    doc.add_heading('3. Chức năng và Nhiệm vụ các phòng ban', level=1)
    
    depts = {
        "Phòng Data Science & AI": [
            "Nghiên cứu và triển khai các mô hình Machine Learning.",
            "Xây dựng mô hình dự báo và tối ưu hóa quy trình cho khách hàng."
        ],
        "Phòng Kỹ thuật Dữ liệu": [
            "Thu thập, làm sạch và chuẩn hóa dữ liệu (ETL).",
            "Đảm bảo hạ tầng server và bảo mật dữ liệu."
        ],
        "Phòng Kinh doanh & Marketing": [
            "Xây dựng thương hiệu và tìm kiếm khách hàng doanh nghiệp.",
            "Tư vấn các gói giải pháp BI và AI."
        ]
    }
    
    for name, tasks in depts.items():
        doc.add_heading(name, level=2)
        for task in tasks:
            doc.add_paragraph(task, style='List Bullet')

    # --- PHẦN 4: LOGIC VẬN HÀNH ---
    doc.add_heading('4. Logic quy trình vận hành (Pseudo-code)', level=1)
    code_snippet = (
        "def process_client_request(client_data):\n"
        "    cleaned_data = DataEng_Dept.clean(client_data)\n"
        "    ai_model = DS_Dept.train_model(cleaned_data)\n"
        "    if ai_model.accuracy > 0.85:\n"
        "        Customer_Success.deliver(ai_model)\n"
        "        return 'Project Completed'\n"
        "    else:\n"
        "        DS_Dept.optimize(ai_model)\n"
        "        return 'Optimizing Model'"
    )
    code_para = doc.add_paragraph()
    run = code_para.add_run(code_snippet)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Lưu file
    file_name = "Bao_cao_To_chuc_DataFlow_AI.docx"
    doc.save(file_name)
    print(f"Đã tạo xong file: {file_name}")

if __name__ == "__main__":
    create_startup_report()