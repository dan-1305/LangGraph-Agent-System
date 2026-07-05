import os
from docx import Document

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def fix_ml_doc(original_path, out_path):
    doc = Document(original_path)
    
    # 1. Delete complete sections
    exclude_prefixes = [
        "1.1.3", "1.4.2", "1.4.3",
        "CHƯƠNG 3", "3.1", "3.2", "3.3", "3.4", "3.5",
        "4.3", "4.4", "4.5",
        "5.2.2", "5.2.3", "5.3.2", "5.3.3", "5.4.2", "5.4.3", "5.5.1", "5.5.3"
    ]
    
    current_section = ""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("CHƯƠNG ") or (len(text) > 0 and text[0].isdigit() and "." in text):
            parts = text.split(" ")
            if text.startswith("CHƯƠNG"):
                current_section = parts[0] + " " + parts[1]
            else:
                current_section = parts[0]
                
        # Delete if in excluded section
        exclude = False
        for prefix in exclude_prefixes:
            if current_section.startswith(prefix):
                exclude = True
                break
                
        # Delete specific bullet points
        if text.startswith("Cấp độ Học sâu:") or text.startswith("Cấp độ Đồ án Tốt nghiệp:"):
            exclude = True
        if text.startswith("Mô hình hóa chuỗi thời gian bằng Học sâu (Deep Learning):"):
            exclude = True
        if text.startswith("Trích xuất đặc trưng ngôn ngữ tự nhiên (NLP):"):
            exclude = True
        if text.startswith("Cơ chế dung hợp tín hiệu đa phương thức (Multimodal Fusion):"):
            exclude = True
        if text.startswith("Nội dung kiến trúc hệ thống của Chương 3:"):
            exclude = True
        if text.startswith("Phân tích đa phương thức:"):
            exclude = True
            
        if exclude:
            delete_paragraph(p)
            continue
            
        # Replace keywords in remaining text
        if "Học sâu" in p.text or "LSTM" in p.text or "GRU" in p.text or "PhoBERT" in p.text or "FinBERT" in p.text or "Đa phương thức" in p.text or "Multimodal" in p.text:
            # Reconstruct paragraph text to remove offending parts or rewrite
            # For simplicity, if it's an overview paragraph, we just clear the runs and write a cleaned version.
            full_text = p.text
            if "Đồ án Tốt nghiệp AI (Đa phương thức):" in full_text:
                delete_paragraph(p)
            elif "Học sâu (Deep Learning): Sử dụng mạng nơ-ron" in full_text:
                delete_paragraph(p)
            elif "1.2.1. Mục tiêu tổng quát" in full_text:
                pass # Keep heading
            elif "Hệ thống kết hợp phân tích dữ liệu biến động giá và đánh giá tâm lý tin tức" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Phát triển hệ thống dự báo định lượng tự động ứng dụng Trí tuệ nhân tạo. Hệ thống tập trung phân tích dữ liệu biến động giá bằng Random Forest và SVM nhằm cung cấp các tín hiệu giao dịch khách quan, hỗ trợ ra quyết định và quản trị rủi ro hiệu quả.")
            elif "1.5.2" in full_text or "1.5.3" in full_text:
                # Let's just delete the mention of chapter 3 and 4
                pass
            elif "Nội dung triển khai thực nghiệm của Chương 4:" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Nội dung triển khai thực nghiệm của Chương 4: Hiện thực hóa các thiết kế lý thuyết thông qua hệ thống kịch bản, dữ liệu thực tế giai đoạn 5 năm đối với các mã cổ phiếu đại diện rổ VN30. Đánh giá định lượng hiệu năng phân loại của thuật toán Random Forest, SVM.")
            elif "5.1.1. Tóm tắt quá trình nâng cấp" in full_text:
                pass # Keep heading
            elif "Ví dụ cụ thể: Mô hình sẽ tích hợp mạng LSTM" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Định hình kiến trúc hệ thống: Bài nghiên cứu đã vạch rõ kiến trúc dòng chảy dữ liệu (Data Pipeline) dự kiến. Mô hình tích hợp các thuật toán Học máy cơ bản để xử lý chuỗi số liệu giá đóng cửa.")
            elif "bằng các siêu phẳng phân tách (SVM) hay các cây quyết định (Random Forest/XGBoost)" in full_text:
                pass # This is fine for ML

    doc.save(out_path)

def fix_dl_doc(original_path, out_path):
    doc = Document(original_path)
    
    exclude_prefixes = [
        "1.1.2", "1.4.1", "1.4.3",
        "2.1", "2.2", "2.4", "2.5", # Delete ML theory except 2.3
        "3.2", "3.3", "3.4", # Delete NLP, Multimodal, Backtesting
        "4.2", "4.4", "4.5", # Delete ML experiment and Multimodal experiment
        "5.2.2", "5.2.3", "5.3.2", "5.3.3", "5.4.3" # Delete NLP/Backtesting conclusions
    ]
    
    current_section = ""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("CHƯƠNG ") or (len(text) > 0 and text[0].isdigit() and "." in text):
            parts = text.split(" ")
            if text.startswith("CHƯƠNG"):
                current_section = parts[0] + " " + parts[1]
            else:
                current_section = parts[0]
                
        exclude = False
        for prefix in exclude_prefixes:
            if current_section.startswith(prefix):
                exclude = True
                break
                
        if text.startswith("Cấp độ Học máy:") or text.startswith("Cấp độ Đồ án Tốt nghiệp:"):
            exclude = True
        if text.startswith("Mô hình hóa bằng Học máy Baseline (Machine Learning):"):
            exclude = True
        if text.startswith("Trích xuất đặc trưng ngôn ngữ tự nhiên (NLP):"):
            exclude = True
        if text.startswith("Cơ chế dung hợp tín hiệu đa phương thức (Multimodal Fusion):"):
            exclude = True
            
        if exclude:
            delete_paragraph(p)
            continue
            
        # Clean keywords in DL doc
        if "Logistic Regression" in p.text or "XGBoost" in p.text or "Random Forest" in p.text or "SVM" in p.text or "FinBERT" in p.text or "PhoBERT" in p.text:
            full_text = p.text
            if "Học máy cơ bản: Sử dụng Random Forest, SVM" in full_text:
                delete_paragraph(p)
            elif "Hệ thống kết hợp phân tích dữ liệu biến động giá và đánh giá tâm lý tin tức" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Phát triển hệ thống dự báo định lượng tự động ứng dụng Học Sâu (Deep Learning). Hệ thống tập trung phân tích chuỗi thời gian bằng mạng LSTM/GRU nhằm cung cấp tín hiệu giao dịch khách quan.")
            elif "Nội dung lý thuyết của Chương 2:" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Nội dung lý thuyết của Chương 2: Thiết lập hệ thống cơ sở lý thuyết về tiền xử lý dữ liệu và tạo đặc trưng (Feature Engineering) như Min-Max Scaler, Log Returns, RSI, MACD để chuẩn bị ma trận 3D cho mạng Học Sâu.")
            elif "Nội dung triển khai thực nghiệm của Chương 4:" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Nội dung triển khai thực nghiệm của Chương 4: Bóc tách cấu hình PyTorch, đánh giá hiệu năng dự báo giá đóng cửa của mô hình LSTM.")
            elif "Ví dụ cụ thể: Mô hình sẽ tích hợp mạng LSTM" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Định hình kiến trúc hệ thống: Mô hình tích hợp mạng LSTM xử lý chuỗi số liệu giá đóng cửa 30 phiên liên tiếp bằng cửa sổ trượt để đưa ra dự báo.")
            elif "bằng các siêu phẳng phân tách (SVM) hay các cây quyết định (Random Forest/XGBoost)" in full_text:
                for r in p.runs: r.text = ""
                p.add_run("Việc đưa ra các quyết định dựa trên thuật toán Học Sâu (LSTM/GRU) giúp đảm bảo tính nhất quán tuyệt đối.")

    # We also need to rename "CHƯƠNG 3" to "CHƯƠNG 2" etc if we want, but since we kept 2.3, CHƯƠNG 2 still has content. So the numbering 1, 2, 3, 4, 5 is preserved and logical!
    doc.save(out_path)

if __name__ == "__main__":
    base_dir = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy"
    original = os.path.join(base_dir, "7. Phạm Công Danh và Nguyễn Công Tiến_Dự báo giá cổ phiếu-chỉ số chứng khoán cơ bản.docx")
    out_ml = os.path.join(base_dir, "Do_an_Hoc_may_va_AI.docx")
    out_dl = os.path.join(base_dir, "Do_an_Hoc_sau.docx")
    
    fix_ml_doc(original, out_ml)
    fix_dl_doc(original, out_dl)
    print("Successfully fixed Do_an_Hoc_may_va_AI.docx and Do_an_Hoc_sau.docx")
