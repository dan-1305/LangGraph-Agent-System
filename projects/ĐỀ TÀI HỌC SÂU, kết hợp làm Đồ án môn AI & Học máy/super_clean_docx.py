import os
from docx import Document

def delete_element(el):
    el.getparent().remove(el)

def clean_tables(doc):
    for table in doc.tables:
        delete_element(table._element)

def clean_toc_and_indexes(doc):
    # We want to delete paragraphs that are part of the old TOC and List of Figures/Tables.
    # Usually they are between "MỤC LỤC" and "LỜI CAM ĐOAN", and "DANH MỤC HÌNH ẢNH VÀ BẢNG" and "LỜI MỞ ĐẦU"
    in_toc = False
    in_lof = False
    
    paragraphs_to_delete = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        
        if text == "MỤC LỤC":
            in_toc = True
            paragraphs_to_delete.append(p)
            continue
            
        if text == "LỜI CAM ĐOAN":
            in_toc = False
            
        if text == "DANH MỤC HÌNH ẢNH VÀ BẢNG" or text == "Danh mục hình ảnh" or text == "Danh mục bảng":
            in_lof = True
            paragraphs_to_delete.append(p)
            continue
            
        if text == "LỜI MỞ ĐẦU":
            in_lof = False
            
        if in_toc or in_lof:
            paragraphs_to_delete.append(p)
            
    for p in paragraphs_to_delete:
        try:
            delete_element(p._element)
        except Exception:
            pass

def replace_text_in_p(p, old, new):
    if old in p.text:
        # A simple but destructive way to replace text in a paragraph while keeping it as one run
        # is to clear runs and set text, but we lose inline bold/italics.
        # Since these are mostly normal text paragraphs, it's acceptable.
        full_text = p.text.replace(old, new)
        for r in p.runs:
            r.text = ""
        p.add_run(full_text)

def clean_dl_doc(path):
    doc = Document(path)
    clean_tables(doc)
    clean_toc_and_indexes(doc)
    
    # Specific DL fixes:
    paragraphs_to_delete = []
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Bị trùng tiêu đề Chương 2: Xóa dòng có "Học máy"
        if "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ KỸ THUẬT HỌC MÁY" in text:
            paragraphs_to_delete.append(p)
            continue
            
        # Tóm tắt chương bị sai lệch
        if "Hồi quy Logistic, Máy học véc-tơ hỗ trợ (SVM), Rừng ngẫu nhiên (Random Forest)" in text:
            replace_text_in_p(p, 
                "Hồi quy Logistic, Máy học véc-tơ hỗ trợ (SVM), Rừng ngẫu nhiên (Random Forest) trong quy trình xếp hạng tầm quan trọng của các tính năng.", 
                "Cơ chế tiền xử lý dữ liệu và tạo tensor đầu vào.")
                
        # Câu chữ dư thừa: "Đây là điểm sáng trong Đồ án AI của bạn"
        if "Đây là điểm sáng trong Đồ án AI của bạn." in text:
            replace_text_in_p(p, 
                "Đây là điểm sáng trong Đồ án AI của bạn.", 
                "Đây là một phương pháp tiếp cận tiên tiến trong nghiên cứu.")
                
    for p in paragraphs_to_delete:
        try:
            delete_element(p._element)
        except Exception:
            pass
            
    doc.save(path)

def clean_ml_doc(path):
    doc = Document(path)
    clean_tables(doc)
    clean_toc_and_indexes(doc)
    
    paragraphs_to_delete = []
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Mục 1.1.3 tâng bốc mạng học sâu, FinBERT
        if "mạng học sâu (LSTM/GRU)" in text or "FinBERT/PhoBERT" in text or "Mô hình ngôn ngữ lớn (LLM" in text:
            # We will just delete this paragraph or rewrite it.
            # In 1.1.3 "Phân tích đa phương thức", we probably shouldn't even have this in ML doc.
            if "Phân tích đa phương thức:" in text:
                paragraphs_to_delete.append(p)
                continue
                
        # Ở mục 1.5.2 tóm tắt chương cũng vẫn nhắc đến "mạng LSTM/GRU... FinBERT"
        if "LSTM/GRU" in text and "FinBERT" in text and "1.5.2" in text:
            paragraphs_to_delete.append(p)
            continue
        if "mạng LSTM/GRU" in text and "FinBERT" in text:
             replace_text_in_p(p,
                 "giải quyết các bài toán học sâu chuỗi thời gian thông qua mạng LSTM/GRU, kỹ thuật trích xuất đặc trưng văn bản tin tức phi cấu trúc dựa trên mô hình ngôn ngữ lớn chuyên biệt FinBERT",
                 "tập trung xử lý bài toán phân loại thông qua các thuật toán Random Forest, SVM")
             replace_text_in_p(p,
                 "kiến trúc dung hợp tensor đặc trưng (Data Fusion), và thiết lập các module kiểm thử chiến lược giao dịch tự động (Backtesting Engine) tích hợp rào cản chi phí thực tế",
                 "thiết lập quy trình tiền xử lý, rút trích đặc trưng kỹ thuật và đánh giá hiệu năng mô hình một cách toàn diện")
                
        # Lỗi chí mạng ở Chương Kết Luận (Mục 4.1.1 hoặc 5.1.1):
        if "Mô hình tích hợp mạng LSTM xử lý chuỗi số liệu giá đóng cửa 30 phiên liên tiếp" in text:
            replace_text_in_p(p,
                "Mô hình tích hợp mạng LSTM xử lý chuỗi số liệu giá đóng cửa 30 phiên liên tiếp",
                "Mô hình tích hợp thuật toán Rừng ngẫu nhiên (Random Forest) xử lý ma trận các chỉ báo động lượng")
                
        if "mạng LSTM xử lý chuỗi số liệu" in text:
            replace_text_in_p(p,
                "mạng LSTM xử lý chuỗi số liệu giá đóng cửa 10 phiên liên tiếp kết hợp song song với mô hình FinBERT để đọc hiểu dữ liệu văn bản từ các tin tức tài chính",
                "thuật toán Rừng ngẫu nhiên (Random Forest) xử lý ma trận các chỉ báo động lượng")

    for p in paragraphs_to_delete:
        try:
            delete_element(p._element)
        except Exception:
            pass
            
    doc.save(path)

if __name__ == "__main__":
    base_dir = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy"
    dl_out = os.path.join(base_dir, "FINAL_Do_an_Hoc_sau.docx")
    ml_out = os.path.join(base_dir, "FINAL_Do_an_Hoc_may_va_AI.docx")
    
    clean_dl_doc(dl_out)
    clean_ml_doc(ml_out)
    print("Super clean completed!")
