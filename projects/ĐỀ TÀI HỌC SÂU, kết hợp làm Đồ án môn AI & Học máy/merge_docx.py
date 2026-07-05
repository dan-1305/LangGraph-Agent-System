import os
import re
from docx import Document
from copy import deepcopy

def delete_paragraphs_after(doc, marker_text):
    found = False
    paragraphs_to_delete = []
    for p in doc.paragraphs:
        if found:
            paragraphs_to_delete.append(p)
        if p.text.strip() == marker_text:
            found = True
    
    for p in paragraphs_to_delete:
        p_el = p._element
        p_el.getparent().remove(p_el)
        p._p = p._element = None

def copy_paragraph(p, doc_dest):
    # We can't directly append a paragraph object to another doc easily while preserving all formatting
    # A workaround is to append a new paragraph and copy text and runs
    new_p = doc_dest.add_paragraph()
    new_p.style = p.style
    new_p.alignment = p.alignment
    for run in p.runs:
        new_run = new_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb
    return new_p

def process_ml(template_path, content_path, output_path):
    doc_out = Document(template_path)
    delete_paragraphs_after(doc_out, "LỜI MỞ ĐẦU")
    
    doc_out.add_paragraph("Thị trường chứng khoán luôn là một môi trường đầu tư hấp dẫn nhưng đầy rủi ro. Đồ án này tập trung vào việc áp dụng Trí tuệ nhân tạo, cụ thể là các thuật toán Học máy cơ bản để dự báo xu hướng thị trường, hỗ trợ ra quyết định đầu tư an toàn và hiệu quả.")
    
    doc_in = Document(content_path)
    
    # Renumbering logic for ML:
    # Chap 1 -> 1
    # Chap 2 -> 2
    # Chap 4 -> 3
    # Chap 5 -> 4
    
    for p in doc_in.paragraphs:
        text = p.text.strip()
        
        # Rewrite CHƯƠNG X
        m_chap = re.match(r'^CHƯƠNG (\d+)(.*)', text)
        if m_chap:
            old_chap = int(m_chap.group(1))
            new_chap = old_chap
            if old_chap == 4: new_chap = 3
            if old_chap == 5: new_chap = 4
            text = f"CHƯƠNG {new_chap}{m_chap.group(2)}"
            
            # Update runs
            p.text = text
            
        else:
            # Rewrite X.Y.Z
            m_sec = re.match(r'^(\d+)\.(\d+)(?:\.(\d+))?(.*)', text)
            if m_sec:
                old_c = int(m_sec.group(1))
                old_s = int(m_sec.group(2))
                old_ss = m_sec.group(3)
                rest = m_sec.group(4)
                
                new_c = old_c
                if old_c == 4: new_c = 3
                if old_c == 5: new_c = 4
                
                new_s = old_s
                # For chapter 4 (now 3), old sections were 4.1, 4.2. We keep 1, 2.
                # For chapter 5 (now 4), old sections were 5.1, 5.2. We keep them as is? Wait, we deleted some sections, so we should renumber the secondary level too if we want perfect sequence, but it's okay to keep the relative numbering if it's contiguous. 
                # Actually, ML Chapter 4 only had 4.1, 4.2 left. So they become 3.1, 3.2.
                # ML Chapter 5 had 5.1, 5.2, 5.3. They become 4.1, 4.2, 4.3.
                
                if old_ss is not None:
                    p.text = f"{new_c}.{new_s}.{old_ss}{rest}"
                else:
                    p.text = f"{new_c}.{new_s}{rest}"

        copy_paragraph(p, doc_out)
        
    doc_out.save(output_path)

def process_dl(template_path, content_path, output_path):
    doc_out = Document(template_path)
    delete_paragraphs_after(doc_out, "LỜI MỞ ĐẦU")
    
    doc_out.add_paragraph("Trong thời đại dữ liệu lớn, việc ứng dụng Học sâu (Deep Learning) vào phân tích chuỗi thời gian tài chính đang mở ra nhiều hướng đi mới. Bài tập lớn này tập trung vào việc xây dựng mô hình mạng nơ-ron hồi quy (LSTM/GRU) nhằm dự báo giá cổ phiếu rổ VN30, tự động hóa quy trình quản trị rủi ro.")
    
    doc_in = Document(content_path)
    
    # Renumbering logic for DL:
    # Chap 1 -> 1
    # Chap 2 -> 2. Section 2.3 -> 2.1
    # Chap 3 -> 3
    # Chap 4 -> 4. Section 4.1 -> 4.1, Section 4.3 -> 4.2
    # Chap 5 -> 5. 
    
    for p in doc_in.paragraphs:
        text = p.text.strip()
        
        m_chap = re.match(r'^CHƯƠNG (\d+)(.*)', text)
        if m_chap:
            old_chap = int(m_chap.group(1))
            # Chapters are 1, 3, 4, 5 in doc_in ? Wait, did I keep CHƯƠNG 2?
            # DL doc has: CHƯƠNG 1, (no CHƯƠNG 2 heading if it was deleted), CHƯƠNG 3, CHƯƠNG 4, CHƯƠNG 5.
            # Actually I need to inject CHƯƠNG 2 heading before 2.3.
            pass # We will handle heading injection below based on 2.3
            
        m_sec = re.match(r'^(\d+)\.(\d+)(?:\.(\d+))?(.*)', text)
        if m_sec:
            old_c = int(m_sec.group(1))
            old_s = int(m_sec.group(2))
            old_ss = m_sec.group(3)
            rest = m_sec.group(4)
            
            new_c = old_c
            new_s = old_s
            
            if old_c == 2:
                # This must be 2.3.x
                if old_s == 3:
                    if old_ss is None:
                        # Before adding 2.1, add CHAPTER 2 heading
                        doc_out.add_paragraph("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ KỸ THUẬT TIỀN XỬ LÝ DỮ LIỆU", style='Heading 1')
                    new_s = 1
            elif old_c == 3:
                new_c = 3
            elif old_c == 4:
                if old_s == 3:
                    new_s = 2 # 4.3 -> 4.2 since 4.2 was deleted
            elif old_c == 5:
                pass
                
            if old_ss is not None:
                p.text = f"{new_c}.{new_s}.{old_ss}{rest}"
            else:
                p.text = f"{new_c}.{new_s}{rest}"

        copy_paragraph(p, doc_out)
        
    doc_out.save(output_path)

if __name__ == "__main__":
    ml_template = "c:/Users/Admin/Downloads/Form Đồ án Tổng hợp Trí tuệ nhân tạo-Học máy.docx"
    dl_template = "c:/Users/Admin/Downloads/Form Bài tập lớn-Học sâu và Ứng dụng.docx"
    
    ml_content = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy/Do_an_Hoc_may_va_AI.docx"
    dl_content = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy/Do_an_Hoc_sau.docx"
    
    ml_out = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy/FINAL_Do_an_Hoc_may_va_AI.docx"
    dl_out = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy/FINAL_Do_an_Hoc_sau.docx"
    
    process_ml(ml_template, ml_content, ml_out)
    process_dl(dl_template, dl_content, dl_out)
    print("Successfully merged into templates!")