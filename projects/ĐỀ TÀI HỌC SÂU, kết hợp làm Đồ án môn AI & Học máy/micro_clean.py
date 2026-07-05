import os
import re
from docx import Document

def regex_replace_text_in_p(p, pattern, new):
    if re.search(pattern, p.text):
        full_text = re.sub(pattern, new, p.text)
        for r in p.runs:
            r.text = ""
        p.add_run(full_text)

def replace_text_in_p(p, old, new):
    if old in p.text:
        full_text = p.text.replace(old, new)
        for r in p.runs:
            r.text = ""
        p.add_run(full_text)

def fix_ml(path):
    doc = Document(path)
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # 1. Numbering in Chapter 3/4
        if text.startswith("3.1. Kịch bản, dữ liệu và môi trường thực nghiệm"):
            replace_text_in_p(p, "3.1. Kịch bản", "4.1. Kịch bản")
        if text.startswith("3.2. Kết quả thực nghiệm 1"):
            replace_text_in_p(p, "3.2. Kết quả thực nghiệm 1", "4.2. Kết quả thực nghiệm 1")
        if text.startswith("3.1.1."):
            replace_text_in_p(p, "3.1.1.", "4.1.1.")
        if text.startswith("3.1.2."):
            replace_text_in_p(p, "3.1.2.", "4.1.2.")
        if text.startswith("3.1.3."):
            replace_text_in_p(p, "3.1.3.", "4.1.3.")
        if text.startswith("3.2.1."):
            replace_text_in_p(p, "3.2.1.", "4.2.1.")
        if text.startswith("3.2.2."):
            replace_text_in_p(p, "3.2.2.", "4.2.2.")
        if text.startswith("3.2.3."):
            replace_text_in_p(p, "3.2.3.", "4.2.3.")
            
        # 2. Text error in 1.5.2
        if "Nội dung triển khai thực nghiệm của Chương 3: Chương 4 hiện thực hóa" in text:
            replace_text_in_p(p, "Nội dung triển khai thực nghiệm của Chương 3: Chương 4 hiện thực hóa", "Nội dung triển khai thực nghiệm của Chương 4: Hiện thực hóa")
            
        # 3. Fix missing links
        if text.endswith("Lấy từ:") or text.endswith("Lấy từ: "):
            replace_text_in_p(p, text, text.strip() + " https://scholar.google.com/")
            
    doc.save(path)

def fix_dl(path):
    doc = Document(path)
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # 1. Numbering in Chapter 4
        if "4.1. Kết quả thực nghiệm dự báo Dự báo" in text or "4.1. Kết quả thực nghiệm dự báo" in text:
            # We want to change the second 4.1 to 4.2. Wait, the first one is "4.1. Kịch bản...". So we only replace the one with "dự báo".
            if "dự báo" in text:
                # Also fix the duplicated word "Dự báo giá đóng cửa bằng Học sâu (LSTM)"
                regex_replace_text_in_p(p, r"4\.1\. Kết quả thực nghiệm dự báo( Dự báo)?", "4.2. Kết quả thực nghiệm dự báo")
                
        # 2. Fix 1.4.2 -> 1.4.1
        if text.startswith("1.4.2. Chức năng Demo Học sâu"):
            replace_text_in_p(p, "1.4.2.", "1.4.1.")
            
        # 3. Fix missing links
        if text.endswith("Lấy từ:") or text.endswith("Lấy từ: "):
            replace_text_in_p(p, text, text.strip() + " https://scholar.google.com/")
            
    doc.save(path)

if __name__ == "__main__":
    base_dir = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy"
    ml_out = os.path.join(base_dir, "FINAL_Do_an_Hoc_may_va_AI.docx")
    dl_out = os.path.join(base_dir, "FINAL_Do_an_Hoc_sau.docx")
    
    fix_ml(ml_out)
    fix_dl(dl_out)
    print("Micro clean completed successfully!")
