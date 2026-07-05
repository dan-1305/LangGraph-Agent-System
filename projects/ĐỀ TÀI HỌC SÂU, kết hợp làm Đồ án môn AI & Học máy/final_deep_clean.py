import os
from docx import Document
import re

def delete_element(el):
    el.getparent().remove(el)

def replace_text_in_p(p, old, new):
    if old in p.text:
        full_text = p.text.replace(old, new)
        for r in p.runs:
            r.text = ""
        p.add_run(full_text)

def regex_replace_text_in_p(p, pattern, new):
    if re.search(pattern, p.text):
        full_text = re.sub(pattern, new, p.text)
        for r in p.runs:
            r.text = ""
        p.add_run(full_text)

def clean_common(doc):
    # 1. Clean TOC and LOF completely
    deleting = False
    paragraphs_to_delete = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        
        if text == "MỤC LỤC":
            deleting = True
            paragraphs_to_delete.append(p)
            continue
            
        if deleting and text == "LỜI CAM ĐOAN":
            deleting = False
            
        if text == "DANH MỤC HÌNH ẢNH VÀ BẢNG" or text == "Danh mục hình ảnh" or text == "Danh mục bảng":
            deleting = True
            paragraphs_to_delete.append(p)
            continue
            
        if deleting and text == "LỜI MỞ ĐẦU":
            deleting = False
            
        if deleting:
            paragraphs_to_delete.append(p)
            
    for p in paragraphs_to_delete:
        try:
            delete_element(p._element)
        except Exception:
            pass

def clean_ml(path):
    doc = Document(path)
    clean_common(doc)
    
    in_refs = False
    paragraphs_to_delete = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # References check
        if text == "DANH MỤC TÀI LIỆU THAM KHẢO":
            in_refs = True
            continue
            
        if in_refs:
            if any(kw in text for kw in ["LSTM", "Deep Learning", "RNN", "BERT", "Hochreiter", "Devlin", "Sequence Modeling"]):
                paragraphs_to_delete.append(p)
                continue
                
        # ML specifics
        if "mạng nơ-ron đa phương thức" in text or "khung làm việc PyTorch" in text or "mạng LSTM và cơ chế Attention" in text:
            regex_replace_text_in_p(p, r"Đối với luồng tính toán nâng cao.*?phần cứng GPU T4\.", "Hệ thống tích hợp môi trường phân tích dữ liệu và các thư viện chuyên dụng như Scikit-learn để xử lý các thuật toán Học máy cơ bản.")
            
        if "mạng LSTM xử lý chuỗi số liệu giá đóng cửa" in text:
            # We want to replace the whole sentence starting with "Mô hình tích hợp mạng LSTM..."
            regex_replace_text_in_p(p, r"Mô hình.*?đóng cửa.*", "Mô hình tích hợp thuật toán Rừng ngẫu nhiên (Random Forest) xử lý ma trận các chỉ báo động lượng.")
            
        if "Nội dung triển khai thực nghiệm của Chương 4" in text:
            replace_text_in_p(p, "Nội dung triển khai thực nghiệm của Chương 4", "Nội dung triển khai thực nghiệm của Chương 3")

    for p in paragraphs_to_delete:
        try:
            delete_element(p._element)
        except Exception:
            pass
            
    doc.save(path)

def clean_dl(path):
    doc = Document(path)
    clean_common(doc)
    
    in_refs = False
    paragraphs_to_delete = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # References check
        if text == "DANH MỤC TÀI LIỆU THAM KHẢO":
            in_refs = True
            continue
            
        if in_refs:
            if any(kw in text for kw in ["Logistic", "SVM", "Random Forest", "XGBoost", "Scikit-learn", "Scikit-Learn", "Machine Learning"]):
                paragraphs_to_delete.append(p)
                continue
                
        # DL specifics
        if "Chương này tập trung vào lý thuyết thị trường, các chỉ báo kỹ thuật tài chính và thuật toán phân loại Học máy để giải quyết bài toán dự đoán xu hướng" in text:
            paragraphs_to_delete.append(p)
            continue
            
        if "(như SVM)" in text or "(như Logistic Regression)" in text:
            replace_text_in_p(p, "(như SVM)", "")
            replace_text_in_p(p, "(như Logistic Regression)", "")
            
        if "Random Forest hay SVM" in text:
            replace_text_in_p(p, "Random Forest hay SVM", "mạng nơ-ron LSTM/GRU")
            
        if "4.2. Kết quả thực nghiệm 2: Dự báo giá đóng cửa bằng Học sâu (LSTM)" in text:
            replace_text_in_p(p, "4.2. Kết quả thực nghiệm 2:", "4.1. Kết quả thực nghiệm dự báo")
            
        if text.startswith("Phân tích đa phương thức:"):
            paragraphs_to_delete.append(p)
            continue
            
        if "Mô hình ngôn ngữ lớn (LLM như FinBERT/PhoBERT)" in text:
            paragraphs_to_delete.append(p)
            continue

    for p in paragraphs_to_delete:
        try:
            delete_element(p._element)
        except Exception:
            pass
            
    doc.save(path)

if __name__ == "__main__":
    base_dir = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy"
    ml_out = os.path.join(base_dir, "FINAL_Do_an_Hoc_may_va_AI.docx")
    dl_out = os.path.join(base_dir, "FINAL_Do_an_Hoc_sau.docx")
    
    clean_ml(ml_out)
    clean_dl(dl_out)
    print("Final deep clean completed successfully!")
