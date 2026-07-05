import os
from docx import Document

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def split_doc(original_path, out_ml_path, out_dl_path):
    # --- Generate ML Doc ---
    doc_ml = Document(original_path)
    
    # Sections to delete for ML doc
    ml_exclude_prefixes = [
        "1.1.3.", "1.3.3.", "1.4.2.", "1.4.3.",
        "CHƯƠNG 3:", "3.1.", "3.1.1.", "3.1.2.", "3.1.3.", "3.2.", "3.2.1.", "3.2.2.", "3.2.3.",
        "3.3.", "3.3.1.", "3.3.2.", "3.3.3.", "3.4.", "3.4.1.", "3.4.2.", "3.4.3.", "3.4.4.",
        "3.5.", "3.5.1.", "3.5.2.", "3.5.3.",
        "4.3.", "4.3.1.", "4.3.2.", "4.3.3.",
        "4.4.", "4.4.1.", "4.4.2.", "4.4.3.",
        "5.2.2.", "5.2.3.", "5.3.2.", "5.3.3.", "5.4.2.", "5.4.3."
    ]
    
    current_section = ""
    for p in doc_ml.paragraphs:
        text = p.text.strip()
        if text.startswith("CHƯƠNG ") or (len(text) > 0 and text[0].isdigit() and "." in text):
            # Update current section if it matches a typical heading
            parts = text.split(" ")
            if text.startswith("CHƯƠNG"):
                current_section = parts[0] + " " + parts[1] # e.g., "CHƯƠNG 3:"
            else:
                current_section = parts[0] # e.g., "1.1.3."
        
        # Check if current section is in excluded prefixes
        exclude = False
        for prefix in ml_exclude_prefixes:
            if current_section.startswith(prefix):
                exclude = True
                break
                
        if exclude:
            delete_paragraph(p)
            
    doc_ml.save(out_ml_path)

    # --- Generate DL Doc ---
    doc_dl = Document(original_path)
    
    dl_exclude_prefixes = [
        "1.2.2. Mục tiêu cụ thể qua từng cấp độ học phần", # This heading has the sub-bullets, we might delete too much, let's be careful.
        # Actually for DL, let's just delete the ML chapters and NLP chapters
        "CHƯƠNG 2:", "2.1.", "2.1.1.", "2.1.2.", "2.1.3.", "2.2.", "2.2.1.", "2.2.2.", "2.2.3.",
        "2.3.", "2.3.1.", "2.3.2.", "2.3.3.", "2.4.", "2.4.1.", "2.4.2.", "2.4.3.", "2.5.", "2.5.1.", "2.5.2.", "2.5.3.",
        "1.3.3.", "1.4.1.", "1.4.3.",
        "3.2.", "3.2.1.", "3.2.2.", "3.2.3.",
        "3.3.", "3.3.1.", "3.3.2.", "3.3.3.",
        "3.4.", "3.4.1.", "3.4.2.", "3.4.3.", "3.4.4.",
        "4.2.", "4.2.1.", "4.2.2.", "4.2.3.",
        "4.4.", "4.4.1.", "4.4.2.", "4.4.3.",
        "5.2.2.", "5.2.3.", "5.3.2.", "5.3.3.", "5.4.2.", "5.4.3."
    ]
    
    current_section = ""
    for p in doc_dl.paragraphs:
        text = p.text.strip()
        if text.startswith("CHƯƠNG ") or (len(text) > 0 and text[0].isdigit() and "." in text):
            parts = text.split(" ")
            if text.startswith("CHƯƠNG"):
                current_section = parts[0] + " " + parts[1]
            else:
                current_section = parts[0]
                
        exclude = False
        for prefix in dl_exclude_prefixes:
            if current_section.startswith(prefix):
                exclude = True
                break
                
        # Handle specific bullet points under 1.2.2 and 1.3.2
        if text.startswith("Cấp độ Học máy:") or text.startswith("Cấp độ Đồ án Tốt nghiệp:"):
            exclude = True
        if text.startswith("Mô hình hóa bằng Học máy Baseline (Machine Learning):") or text.startswith("Trích xuất đặc trưng ngôn ngữ tự nhiên (NLP):"):
            exclude = True
            
        if exclude:
            delete_paragraph(p)
            
    doc_dl.save(out_dl_path)

if __name__ == "__main__":
    base_dir = "projects/ĐỀ TÀI HỌC SÂU, kết hợp làm Đồ án môn AI & Học máy"
    original = os.path.join(base_dir, "7. Phạm Công Danh và Nguyễn Công Tiến_Dự báo giá cổ phiếu-chỉ số chứng khoán cơ bản.docx")
    out_ml = os.path.join(base_dir, "Do_an_Hoc_may_va_AI.docx")
    out_dl = os.path.join(base_dir, "Do_an_Hoc_sau.docx")
    
    split_doc(original, out_ml, out_dl)
    print("Successfully created Do_an_Hoc_may_va_AI.docx and Do_an_Hoc_sau.docx")
