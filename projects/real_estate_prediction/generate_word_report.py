import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Cấu hình Style chung
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # TIÊU ĐỀ BÁO CÁO
    title = doc.add_heading('BÁO CÁO TIỂU LUẬN CHI TIẾT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Tên đề tài: Phân tích dữ liệu và xây dựng mô hình dự báo giá bất động sản tại khu vực TP.HCM và Đồng Nai')
    run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\n')

    # CHƯƠNG 1
    doc.add_heading('CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI', level=1)
    
    doc.add_heading('1.1. Mục tiêu tổng quát', level=2)
    doc.add_paragraph('Mục tiêu tổng quát của đề tài là ứng dụng quy trình Khoa học Dữ liệu (Data Science Pipeline) để xây dựng một hệ thống Trí tuệ Nhân tạo (AI) toàn diện, có khả năng tự động hóa việc định giá bất động sản tại khu vực TP.HCM và Đồng Nai. Hệ thống này hướng tới việc giải quyết bài toán thiếu minh bạch về giá trị nhà đất thông qua việc phân tích lượng lớn dữ liệu thô, từ đó cung cấp một thước đo giá trị khách quan giúp người mua, người bán đưa ra các quyết định tài chính an toàn.')

    doc.add_heading('1.2. Lý do chọn đề tài', level=2)
    doc.add_paragraph('Thị trường bất động sản (BĐS) tại khu vực kinh tế trọng điểm phía Nam (cụ thể là TP.HCM và Đồng Nai) luôn biến động mạnh mẽ. Giá trị BĐS bị chi phối bởi vô số yếu tố: từ định lượng (diện tích, số tầng) đến định tính (vị trí, pháp lý, tiện ích). Đối với bộ môn Nhập môn Khoa học Dữ liệu, đây là một bài toán hoàn hảo để thực hành trọn vẹn quy trình từ thu thập dữ liệu (Data Collection), làm sạch (Data Cleaning), khai phá đặc trưng (Feature Engineering) cho đến mô hình hóa (Modeling).')

    doc.add_heading('1.3. Mục tiêu nghiên cứu cụ thể', level=2)
    doc.add_paragraph('• Thu thập bộ dữ liệu thực tế từ các nền tảng BĐS lớn tập trung vào khu vực TP.HCM và Đồng Nai.')
    doc.add_paragraph('• Thực hiện làm sạch dữ liệu chuyên sâu và loại bỏ các giá trị ngoại lai (Outliers) để đảm bảo chất lượng đầu vào.')
    doc.add_paragraph('• Ứng dụng kỹ thuật khai phá văn bản (Regex) để trích xuất các đặc trưng quan trọng từ tiêu đề tin đăng (mặt tiền, số lầu, tính pháp lý).')
    doc.add_paragraph('• Huấn luyện và đánh giá mô hình học máy (XGBoost) nhằm dự báo giá nhà.')
    doc.add_paragraph('• Xây dựng ứng dụng Web minh họa để đưa mô hình vào thực tiễn.')

    doc.add_heading('1.4. Đối tượng và phạm vi nghiên cứu', level=2)
    doc.add_paragraph('• Đối tượng nghiên cứu: Quy trình phân tích dữ liệu, các phương pháp làm sạch ngoại lai (Z-Score), kỹ thuật Text Mining và thuật toán hồi quy Gradient Boosting.')
    doc.add_paragraph('• Phạm vi nghiên cứu: Dữ liệu tin đăng BĐS tại khu vực TP.Hồ Chí Minh và Tỉnh Đồng Nai, giới hạn phân khúc nhà đất, căn hộ có diện tích từ 10m2 đến 1000m2.')

    # CHƯƠNG 2
    doc.add_heading('CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ TỔNG QUAN TÀI LIỆU', level=1)
    
    doc.add_heading('2.1. Các yếu tố định giá Bất động sản', level=2)
    doc.add_paragraph('Giá trị bất động sản được cấu thành từ:\n- Yếu tố vật lý: Diện tích sàn, số lượng tầng, chiều ngang, và độ rộng hẻm tiếp cận.\n- Yếu tố vị trí & pháp lý: Khu vực hành chính (Quận/Huyện thuộc TP.HCM hoặc Đồng Nai), tình trạng mặt tiền hay hẻm, và sự bảo đảm về mặt pháp lý (Sổ hồng, sổ đỏ).')

    doc.add_heading('2.2. Khai phá dữ liệu văn bản (Text Mining)', level=2)
    doc.add_paragraph('Khác với các bộ dữ liệu chuẩn hóa, dữ liệu BĐS trên mạng phần lớn là văn bản tự do. Biểu thức chính quy (Regex) được sử dụng như một kỹ thuật cốt lõi trong Data Science để bóc tách thông tin ẩn (vd: "nhà 3 lầu, hẻm 4m, SHR") thành các biến số phân loại độc lập có thể tính toán được.')

    doc.add_heading('2.3. Thuật toán XGBoost trong Hồi quy', level=2)
    doc.add_paragraph('XGBoost (eXtreme Gradient Boosting) là thuật toán Ensemble Learning tiên tiến. Khả năng chống Overfitting mạnh mẽ và khả năng xử lý tốt cả biến phân loại lẫn biến liên tục giúp XGBoost vượt trội trong bài toán định giá BĐS.')

    # CHƯƠNG 3
    doc.add_heading('CHƯƠNG 3: PHƯƠNG PHÁP NGHIÊN CỨU', level=1)
    doc.add_paragraph('Nghiên cứu được thiết kế bám sát vào Quy trình Khoa học Dữ liệu (Data Science Pipeline) bao gồm 5 bước cốt lõi:\n1. Data Collection: Thu thập dữ liệu thô.\n2. Data Cleaning & EDA: Xử lý dữ liệu khuyết thiếu và phân tích khám phá để hiểu phân phối giá.\n3. Feature Engineering: Làm giàu dữ liệu thông qua việc biến đổi biến mục tiêu (Log Transformation) và trích xuất đặc trưng mới.\n4. Modeling: Xây dựng Pipeline huấn luyện với tỷ lệ chia tập dữ liệu 80:20.\n5. Deployment: Triển khai mô hình dưới dạng dịch vụ Web.')

    # CHƯƠNG 4
    doc.add_heading('CHƯƠNG 4: THU THẬP VÀ XỬ LÝ DỮ LIỆU', level=1)
    
    doc.add_heading('4.1. Thu thập dữ liệu', level=2)
    doc.add_paragraph('Dữ liệu đầu vào (raw_data.csv) được cào tự động bằng thư viện Playwright từ các nền tảng bất động sản lớn, với thiết lập vị trí đích danh là TP.HCM và Đồng Nai. Cấu trúc dữ liệu ban đầu bao gồm: Giá (Price_VND), Diện tích (Area_m2), Quận/Huyện (District), Tỉnh (Province), và Tiêu đề (Title).')

    doc.add_heading('4.2. Tiền xử lý dữ liệu (Data Preprocessing)', level=2)
    doc.add_paragraph('Quá trình làm sạch dữ liệu là khâu quan trọng nhất quyết định độ chính xác của toàn bộ dự án:\n- Xử lý missing values: Loại bỏ lập tức các dòng bị khuyết thông tin trọng yếu như Giá và Diện tích.\n- Lọc tính hợp lý lý thuyết: Chỉ giữ lại các bản ghi có Area_m2 từ 10m2 đến 1000m2.\n- Lọc biên độ giá: Tạo ra biến Đơn giá (Price_per_m2). Áp dụng ngưỡng cứng: loại bỏ bất động sản có đơn giá bất hợp lý (< 1 triệu hoặc > 300 triệu VNĐ/m2) nhằm loại bỏ rác/tin spam.')

    doc.add_heading('4.3. Xử lý dữ liệu ngoại lai (Outlier Detection) bằng Z-Score', level=2)
    doc.add_paragraph('Do giá nhà tại TP.HCM khác biệt rất lớn so với Đồng Nai, phương pháp Z-Score cục bộ được áp dụng thay vì Z-Score toàn cục:\n- Hệ thống tự động tính Giá trị trung bình (Mean) và Độ lệch chuẩn (Std) của Đơn giá theo từng khu vực Quận/Huyện cụ thể.\n- Tính điểm Z-Score cho mỗi bản ghi. Dữ liệu có |Z| > 3 (lệch quá 3 độ lệch chuẩn) bị xem là giá "ảo" và bị loại bỏ. Khâu này giúp dữ liệu trở nên cực kỳ sạch và đáng tin cậy.')

    doc.add_heading('4.4. Trích xuất đặc trưng (Feature Engineering)', level=2)
    doc.add_paragraph('Nhằm tối đa hóa thông tin từ cột Title, kỹ thuật xử lý ngôn ngữ cơ bản được áp dụng để sinh ra 8 đặc trưng mới, tạo sức mạnh cho mô hình máy học:\n1. is_frontage: Nhận diện "mặt tiền", "mặt đường".\n2. is_alley: Nhận diện "hẻm", "hxh".\n3. has_furniture: Tình trạng nội thất ("full đồ").\n4. num_floors: Số lượng lầu/tầng (dạng số).\n5. width_m: Bề ngang lô đất.\n6. alley_width_m: Độ rộng hẻm tiếp cận.\n7. Property_Type: Phân loại thành Nhà, Đất, Căn hộ, Kho xưởng.\n8. is_so_hong: Đảm bảo tính pháp lý.')

    # CHƯƠNG 5
    doc.add_heading('CHƯƠNG 5: KẾT QUẢ NGHIÊN CỨU', level=1)
    
    doc.add_heading('5.1. Thiết lập huấn luyện mô hình', level=2)
    doc.add_paragraph('Do biến mục tiêu (Price_VND) có phân phối lệch rất mạnh, dữ liệu được chuyển đổi qua hàm Logarit tự nhiên (np.log1p) để đưa về phân phối chuẩn (Normal Distribution). Điểm nổi bật trong kiến trúc mới là Hệ thống Phân mảnh (Segmentation Logic). Dữ liệu được chia làm 2 phân khúc: Cao Cấp (High-End) và Phổ Thông (Mass-Market). Mỗi phân khúc được huấn luyện bằng một mô hình XGBoost riêng biệt.')
    doc.add_paragraph('Đồng thời, hệ thống ứng dụng RandomizedSearchCV kết hợp K-Fold Cross Validation để tự động dò tìm siêu tham số (Hyperparameter Tuning), giúp chống Overfitting hiệu quả. Dữ liệu trước khi vào mô hình được xử lý chuẩn hóa bằng StandardScaler và mã hóa vector qua OneHotEncoder.')

    doc.add_heading('5.2. Kết quả đánh giá mô hình', level=2)
    doc.add_paragraph('Trên tập kiểm thử (Test set), các mô hình đạt được kết quả khả quan và kiểm soát tốt hiện tượng Overfitting:\n- Phân khúc High-End: Chỉ số R-squared (R^2) đạt ~0.4525. Sai số tuyệt đối trung bình (MAE) xấp xỉ ~4.6 Tỷ VNĐ.\n- Phân khúc Mass-Market: Chỉ số R-squared (R^2) đạt ~0.3731. Sai số tuyệt đối trung bình (MAE) xấp xỉ ~1.93 Tỷ VNĐ.\nViệc tách phân khúc và tối ưu tham số giúp hệ thống dự báo bám sát hơn với đặc thù và mức dao động giá trị rất khác biệt của thị trường bất động sản.')

    doc.add_heading('5.3. Triển khai ứng dụng (Deployment)', level=2)
    doc.add_paragraph('Thể hiện đúng tinh thần của khoa học dữ liệu là mang lại giá trị thực tiễn, nhóm đã xuất mô hình thành file nhị phân (model.pkl) và xây dựng ứng dụng Web với Flask. Giao diện trực quan cho phép người dùng cấu hình tham số BĐS và nhận lại mức giá tham khảo ngay tức thì.')

    # CHƯƠNG 6
    doc.add_heading('CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', level=1)
    
    doc.add_heading('6.1. Kết luận', level=2)
    doc.add_paragraph('Đề tài đã hoàn thành một vòng đời (life-cycle) hoàn chỉnh của dự án Khoa học Dữ liệu: từ cào dữ liệu, làm sạch sâu, xử lý outlier, trích xuất đặc trưng văn bản, cho đến mô hình hóa và triển khai thành công ứng dụng thực tiễn cho thị trường TP.HCM và Đồng Nai.')

    doc.add_heading('6.2. Hạn chế và Hướng phát triển', level=2)
    doc.add_paragraph('- Hạn chế: Các luật Regex phân tích văn bản tĩnh chưa thể thấu hiểu hoàn toàn các từ lóng hay cách hành văn viết tắt phức tạp của môi giới BĐS.\n- Phát triển: Trong các nghiên cứu nâng cao hơn, có thể tích hợp mô hình ngôn ngữ như PhoBERT để xử lý triệt để ngữ nghĩa văn bản. Đồng thời, thu thập thêm dữ liệu về khoảng cách địa lý (tọa độ GPS, khoảng cách đến trường học/bệnh viện) thông qua OpenStreetMap để tăng độ chính xác (R^2) của mô hình định giá.')

    # Lưu file
    output_path = os.path.join("projects", "real_estate_prediction", "reports", "Bao_Cao_Tieu_Luan_BDS.docx")
    doc.save(output_path)
    print(f"Đã tạo thành công file: {output_path}")

if __name__ == "__main__":
    create_report()