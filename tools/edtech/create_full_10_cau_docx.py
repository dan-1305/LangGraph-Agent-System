from docx import Document
from docx.shared import Pt, RGBColor

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('HƯỚNG DẪN GIẢI CHI TIẾT - ĐỀ CƯƠNG BMTT (FULL 10 CÂU)', 0)
    doc.add_paragraph('Tài liệu chép tay, bao gồm tư duy (Mindset) màu Xanh dương và nội dung cần chép vào giấy thi màu Xanh lá.')

    def add_mindset(text):
        p = doc.add_paragraph()
        run = p.add_run('💭 Mindset (Nháp trong đầu):')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 112, 192)
        doc.add_paragraph(text)

    def add_write(text):
        p = doc.add_paragraph()
        run = p.add_run('✍️ Viết ra giấy thi:')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 176, 80)
        doc.add_paragraph(text)

    # CÂU 1
    doc.add_heading('Câu 1: Sử dụng Mã Affine giải mã bản mã "rmpjvpwtwtmhgmpwtgfwgkzn" với a=19, b=23', level=1)
    add_mindset('- Công thức giải mã: P = a^(-1) * (C - b) mod 26.\n- Tìm nghịch đảo của 19 theo mod 26: 19 * 11 = 209 ≡ 1 mod 26 => a^(-1) = 11.\n- (C - 23) ≡ (C + 3) mod 26. Vậy P = 11 * (C + 3) mod 26.')
    add_write('- Công thức giải mã Affine: P = a^(-1) * (C - b) mod 26.\n- Ta có 19 * 11 ≡ 1 mod 26 => a^(-1) = 11.\n- Suy ra công thức: P = 11 * (C - 23) ≡ 11 * (C + 3) mod 26.\n- Ráp C=17 (r) -> P = 11*(17+3) mod 26 = 12 (M).\n- Ráp C=12 (m) -> P = 11*(12+3) mod 26 = 9 (J).\n- (Kẻ bảng tính toán tương tự cho các ký tự còn lại)\n=> Bản rõ: MJQCEQPIPIJGVJQPIVKPVNWU')

    # CÂU 2
    doc.add_heading('Câu 2: Mã hóa & giải mã Hill "doikhongnhulamo" với khóa K = [[7, 3], [8, 7]]', level=1)
    add_mindset('- Hill 2x2. Chuỗi có 15 ký tự, lẻ -> Thêm "x" vào cuối thành 16 ký tự, 8 block 2 ký tự: do, ik, ho, ng, nh, ul, am, ox.\n- TRỌNG TÂM DỄ SAI: Mã hóa Hill tính theo dạng Vector CỘT. C = K x P mod 26. Tức là C1 = K11*P1 + K12*P2.\n- Block "do"=[3, 14]: C1 = 7*3 + 3*14 = 63 ≡ 11 (L), C2 = 8*3 + 7*14 = 122 ≡ 18 (S) => LS.\n- Giải mã: C x K^(-1) mod 26. Với K^(-1) = det(K)^(-1) * Adj(K) mod 26.')
    add_write('A. MÃ HÓA\n- Bản rõ (thêm x ở cuối): do ik ho ng nh ul am ox -> (3,14), (8,10), (7,14), (13,6), (13,7), (20,11), (0,12), (14,23).\n- C = K x P mod 26. Block 1: [[7,3], [8,7]] x [3, 14]^T = [63, 122]^T ≡ [11, 18]^T mod 26 => LS.\n- Block 2 (ik): [[7,3], [8,7]] x [8, 10]^T = [86, 134]^T ≡ [8, 4]^T mod 26 => IE.\n- Làm tương tự: LSIENYFQIXRDKGLN.\nB. GIẢI MÃ\n- det(K) = 7*7 - 3*8 = 25 ≡ -1 mod 26. Suy ra det(K)^(-1) = 25 mod 26.\n- Adj(K) = [[7, -3], [-8, 7]] ≡ [[7, 23], [18, 7]] mod 26.\n- K^(-1) = 25 * [[7, 23], [18, 7]] ≡ [[19, 3], [8, 19]] mod 26.\n- P = K^(-1) x C mod 26. Block 1: [[19,3], [8,19]] x [11, 18]^T = [263, 430]^T ≡ [3, 14]^T mod 26 => do.')

    # CÂU 3
    doc.add_heading('Câu 3: Sử dụng Mã Hill mã hóa "khoacongnghekythuat" với K = [[5, 8, 9], [3, 7, 5], [6, 3, 9]]', level=1)
    add_mindset('- Hill 3x3. Chuỗi "khoacongnghekythuat" có 19 ký tự. Thêm 2 "x" thành 21 ký tự (7 block 3 ký tự): kho, aco, ngn, ghe, kyt, hua, txx.\n- Mã hóa: C = K x P mod 26. (Ma trận 3x3 nhân vector cột 3x1).\n- Giải mã: Tính det(K) = 19 => det(K)^(-1) = 11. Tính ma trận phụ hợp bằng phần bù đại số.\n- K^(-1) sẽ là: [[8, 25, 7], [7, 5, 22], [1, 25, 17]].')
    add_write('A. MÃ HÓA\n- Bản rõ (thêm xx): kho aco ngn ghe kyt hua txx -> (10,7,14), (0,2,14), (13,6,13), (6,7,4), (10,24,19), (7,20,0), (19,23,23).\n- C = K x P mod 26. Block 1 (kho): K x [10,7,14]^T = [232, 149, 207]^T ≡ [24, 19, 25]^T mod 26 => YTZ.\n- Kẻ bảng tính tiếp tục cho các block còn lại. Kết quả mã hóa: YTZMGCWQFSJPXHRNFYSVA\nB. GIẢI MÃ\n- det(K) = 5(63-15) - 8(27-30) + 9(9-42) = 240 + 24 - 297 = -33 ≡ 19 mod 26.\n- det(K)^(-1) ≡ 11 mod 26.\n- Tính ma trận phần bù đại số và chuyển vị để lấy Adj(K). Suy ra K^(-1) = 11 * Adj(K) mod 26 = [[8, 25, 7], [7, 5, 22], [1, 25, 17]].\n- P = K^(-1) x C mod 26. Block 1: K^(-1) x [24, 19, 25]^T = [842, 813, 924]^T ≡ [10, 7, 14]^T mod 26 => kho.')

    # CÂU 4
    doc.add_heading('Câu 4: Mã hóa RSA nội dung "chuccacbanthitotnhe" với p=11, q=13', level=1)
    add_mindset('- Tính n = 11*13=143; φ(n) = 10*12=120.\n- Chọn e=7 (vì gcd(7,120)=1). Khóa PU = (7, 143).\n- Tìm d sao cho 7*d = 1 mod 120 => d=103. Khóa PR = (103, 143).\n- Mã hóa C = M^e mod n. Ví dụ "c"=2 -> 2^7 mod 143 = 128. "h"=7 -> 7^7 mod 143 = 6.')
    add_write('- Bước 1: Tính n = p*q = 143. Tính φ(n) = (p-1)(q-1) = 120.\n- Bước 2: Chọn e = 7. Khóa công khai PU = (7, 143).\n- Bước 3: Tính d ≡ 7^(-1) mod 120 => d = 103. Khóa bí mật PR = (103, 143).\n- Bước 4: Mã hóa C = M^e mod 143.\nChữ c (2) => C1 = 2^7 mod 143 = 128.\nChữ h (7) => C2 = 7^7 mod 143 = 6.\nChữ u (20) => C3 = 20^7 mod 143 = 136.\nChữ c (2) => C4 = 2^7 mod 143 = 128.\nChữ c (2) => C5 = 128.\n(Ghi tiếp 1 vài ký tự đại diện). Chuỗi số mã hóa: [128, 6, 136, 128, 128, 0, 128, 1, 0, 117, 46, 6, 57, 46, 53, 46, 117, 6, 82].')

    # CÂU 5
    doc.add_heading('Câu 5: Hãy trình Hệ thống ngăn ngừa xâm nhập IPS', level=1)
    add_write('- IPS (Intrusion Prevention System) là hệ thống bảo mật chủ động.\n- Vai trò: Nằm trực tiếp trên luồng dữ liệu (Inline). Có khả năng phát hiện các hành vi có hại và tự động chặn đứng, ngắt kết nối (Drop/Block) cuộc tấn công ngay lập tức.\n- Cơ chế hoạt động: Dựa trên chữ ký (Signature-based) hoặc sự bất thường (Anomaly-based).')

    # CÂU 6
    doc.add_heading('Câu 6: Mã hóa "họ và tên" của bạn bằng Playfair với key="bainaylamroi"', level=1)
    add_mindset('- Lập bảng 5x5 từ key, bỏ chữ trùng, gộp I/J.\n- B A I N Y | L M R O C | D E F G H | K P Q S T | U V W X Z.')
    add_write('Bước 1: Bảng 5x5: B-A-I-N-Y | L-M-R-O-C | D-E-F-G-H | K-P-Q-S-T | U-V-W-X-Z.\nBước 2: Tên (Ví dụ) NGUYEN VAN A. Tách block: NG UY EN VA NA.\nBước 3: Mã hóa.\n- NG (Hình chữ nhật) -> YF.\n- UY (Hình chữ nhật) -> ZB.\n- EN (Hình chữ nhật) -> HA.\n- VA (Cùng cột) -> AV.\n- NA (Cùng hàng) -> YI.\nBản mã: YF ZB HA AV YI.')

    # CÂU 7
    doc.add_heading('Câu 7: Mã hóa họ và tên của bạn bằng Playfair với key="lamnhieuroi"', level=1)
    add_write('Bước 1: Bảng 5x5 từ khóa "lamnhieuroi": L-A-M-N-H | I-E-U-R-O | B-C-D-F-G | K-P-Q-S-T | V-W-X-Y-Z.\nBước 2: Tên (Ví dụ) NGUYEN VAN A. Tách block: NG UY EN VA NA.\nBước 3: Mã hóa (tương tự quy tắc cùng hàng, cùng cột, hình chữ nhật).\nBản mã: FH RZ HU LW LM.')

    # CÂU 8
    doc.add_heading('Câu 8: Hãy trình bày sự khác nhau giữa IPS và IDS', level=1)
    add_write('- Tính chất: IDS (Phát hiện) là thụ động, IPS (Ngăn ngừa) là chủ động.\n- Vị trí: IDS thường đứng ngoài luồng mạng quan sát (Out-of-band), IPS nằm trực tiếp trong luồng mạng (Inline).\n- Phản ứng: IDS chỉ đưa ra cảnh báo (Alert) cho quản trị viên. IPS có khả năng tự động ngắt kết nối và chặn cuộc tấn công.')

    # CÂU 9
    doc.add_heading('Câu 9: Trình bày vai trò của SSL/TLS trong bảo mật website', level=1)
    add_write('1. Mã hóa (Encryption): Mã hóa dữ liệu trên đường truyền, ngăn chặn kẻ tấn công đọc trộm thông tin (ví dụ: mật khẩu, thẻ tín dụng).\n2. Xác thực (Authentication): Sử dụng chứng chỉ số (Certificate) để chứng minh danh tính của máy chủ với người dùng.\n3. Tính toàn vẹn (Integrity): Đảm bảo dữ liệu không bị kẻ gian thay đổi trong quá trình truyền đi.')

    # CÂU 10
    doc.add_heading('Câu 10: Trình bày cơ chế hoạt động của chữ ký số', level=1)
    add_write('Cơ chế hoạt động gồm 2 quá trình dựa trên mã hóa khóa công khai:\n1. Quá trình Ký (Người gửi): Người gửi dùng hàm băm (Hash) tóm tắt tài liệu, sau đó dùng Khóa bí mật (Private Key) để mã hóa bản tóm tắt tạo thành Chữ ký số đính kèm.\n2. Quá trình Kiểm tra (Người nhận): Người nhận dùng Khóa công khai (Public Key) của người gửi để giải mã chữ ký, thu lại bản tóm tắt gốc. Sau đó so sánh với bản tóm tắt mới tự tính để xác thực.\nVai trò: Chống giả mạo, toàn vẹn dữ liệu, chống chối bỏ.')

    doc.save('docs/BMTT/Giai_Viet_Tay_Full_10_Cau.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()