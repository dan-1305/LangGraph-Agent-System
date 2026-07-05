from docx import Document
from docx.shared import Pt

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('BẢNG TRA CỨU NHANH BMTT (MANG VÀO PHÒNG THI)', 0)
    doc.add_paragraph('Bảng chuyển đổi chữ cái tiếng Anh sang số (Dùng cho Affine, Hill, Vigenere...)')
    
    table = doc.add_table(rows=2, cols=13)
    table.style = 'Table Grid'
    
    row0 = table.rows[0].cells
    row1 = table.rows[1].cells
    for i in range(13):
        row0[i].text = chr(65 + i) # A-M
        row1[i].text = str(i)
        
    doc.add_paragraph() # space
    
    table2 = doc.add_table(rows=2, cols=13)
    table2.style = 'Table Grid'
    
    row0_2 = table2.rows[0].cells
    row1_2 = table2.rows[1].cells
    for i in range(13, 26):
        row0_2[i-13].text = chr(65 + i) # N-Z
        row1_2[i-13].text = str(i)

    doc.add_paragraph('\nLưu ý:')
    doc.add_paragraph('- Mọi tính toán đều thực hiện trên Modulo 26 (chỉ dùng các số từ 0 đến 25).')
    doc.add_paragraph('- TUYỆT ĐỐI KHÔNG sử dụng bảng mã ASCII cho các bài tập này vì các bài Affine, Hill, Playfair đều quy ước A=0, B=1... Z=25.')
    doc.add_paragraph('- Mã Playfair sử dụng bảng 5x5 nên sẽ gộp chữ I và chữ J làm một (coi như giống nhau).')

    doc.save('docs/BMTT/Bang_Tra_Cuu_BMTT.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()