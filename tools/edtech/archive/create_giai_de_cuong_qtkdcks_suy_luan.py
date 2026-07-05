import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = None

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('SỔ TAY SUY LUẬN BÀI TẬP QTKD (HƯỚNG DẪN TƯ DUY)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True

    add_heading(doc, "CÂU 5: BÀI TOÁN TỒN KHO CƠ BẢN (EOQ)", level=1)
    
    add_paragraph(doc, "1. Dịch Đề (Dữ liệu đã cho):", bold=True)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Thấy chữ \"Nhu cầu hằng năm\" => Ký hiệu là ").bold = False
    p.add_run("D").bold = True
    p.add_run(" (Demand).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Thấy chữ \"Chi phí đặt hàng mỗi lần\" => Ký hiệu là ").bold = False
    p.add_run("S").bold = True
    p.add_run(" (Setup cost).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Thấy chữ \"Chi phí tồn trữ (lưu kho)\" => Ký hiệu là ").bold = False
    p.add_run("H").bold = True
    p.add_run(" (Holding cost).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Thấy \"Thời gian giao hàng\" => Ký hiệu là ").bold = False
    p.add_run("L").bold = True
    p.add_run(" (Lead time).")
    
    add_paragraph(doc, "2. Suy luận giải (Hỏi gì tính nấy):", bold=True)
    add_paragraph(doc, "- Đề yêu cầu tính Lượng đặt hàng kinh tế (EOQ):", bold=True)
    add_paragraph(doc, "  => Lôi công thức: EOQ = √(2.D.S / H)", italic=True)
    add_paragraph(doc, "- Đề yêu cầu tính Số lần đặt hàng (N):", bold=True)
    add_paragraph(doc, "  => Tư duy: Một năm cần D cái, mỗi lần đặt EOQ cái. Số lần = Tổng chia 1 lần.")
    add_paragraph(doc, "  => Công thức: N = D / EOQ", italic=True)
    add_paragraph(doc, "- Đề yêu cầu tính Thời gian giữa 2 lần đặt (Chu kỳ - T):", bold=True)
    add_paragraph(doc, "  => Tư duy: Một năm làm việc bao nhiêu ngày, chia đều cho số lần đặt.")
    add_paragraph(doc, "  => Công thức: T = Số ngày làm việc / N", italic=True)
    add_paragraph(doc, "- Đề yêu cầu tính Tổng chi phí (TC):", bold=True)
    add_paragraph(doc, "  => Tư duy: TC = Phí Đặt + Phí Tồn. (Không tính phí mua vì giá không đổi).")
    add_paragraph(doc, "  => Công thức: TC = (D/EOQ).S + (EOQ/2).H", italic=True)
    add_paragraph(doc, "- Đề yêu cầu tính Điểm tái đặt hàng (ROP):", bold=True)
    add_paragraph(doc, "  => Tư duy: ROP là lúc trong kho còn bao nhiêu thì phải gọi hàng. Cần tìm 1 ngày bán được bao nhiêu (d), nhân với số ngày chờ hàng (L), cộng với dự trữ an toàn (nếu có).")
    add_paragraph(doc, "  => Công thức: d = D / Số ngày làm việc. Rồi tính ROP = d × L + SS", italic=True)
    
    add_paragraph(doc, "⚠️ LƯU Ý CHẾT NGƯỜI (TRÁNH MẤT ĐIỂM):", bold=True, color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- Kiểm tra xem đề cho chi phí tồn trữ H là SỐ TIỀN hay TỶ LỆ (%). Nếu cho %. => H = i.P (Tỷ lệ × Giá mua).", color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- Lượng EOQ tính ra lẻ thì LÀM TRÒN LÊN/XUỐNG tùy bạn nhưng nhớ ghi ≈ (VD: 1.549 vòng bi, không được ghi 1.549,2 vòng bi vì nó là vật thể).", color=RGBColor(255, 0, 0))

    add_heading(doc, "CÂU 6: BÀI TOÁN TỒN KHO CHIẾT KHẤU", level=1)
    
    add_paragraph(doc, "1. Dịch Đề (Khác bài cơ bản chỗ nào?):", bold=True)
    add_paragraph(doc, "- Đề sẽ cho một cái bảng giá: Mua từ X đến Y giá là P1, mua lớn hơn Z giá là P2... Mua càng nhiều giá càng rẻ.")
    add_paragraph(doc, "- Chi phí tồn trữ (H) thường sẽ cho dạng tỷ lệ (VD: 15% giá mua). Do giá mua P thay đổi nên H cũng thay đổi theo từng mức giá.")

    add_paragraph(doc, "2. Suy luận giải (Giải làm 4 bước):", bold=True)
    add_paragraph(doc, "Bước 1: Tính riêng lẻ EOQ cho TỪNG MỨC GIÁ (từ thấp nhất đến cao nhất).")
    add_paragraph(doc, "  => Nhớ là ở mức giá P nào thì tính bằng công thức: EOQi = √(2.D.S / (i.Pi))")
    add_paragraph(doc, "Bước 2: Xét tính HỢP LỆ của EOQ (Kiểm tra xem EOQ vừa tính có nằm trong khoảng điều kiện của cái giá đó không).")
    add_paragraph(doc, "  => Nếu hợp lệ (Nằm trong khoảng): Giữ nguyên làm ứng cử viên Q*.")
    add_paragraph(doc, "  => Nếu không hợp lệ (Bé hơn mức quy định để được giá đó): Đẩy nó lên bằng mức tối thiểu của khoảng chiết khấu đó.")
    add_paragraph(doc, "Bước 3: Đưa các Q* (ứng cử viên) vừa chốt ở Bước 2 vào tính TỔNG CHI PHÍ TC.")
    add_paragraph(doc, "  => Nhớ: Ở bài này BẮT BUỘC phải cộng cả Giá mua vào TC vì giá mua bị thay đổi giữa các phương án.")
    add_paragraph(doc, "  => TC = D.P + (D/Q*).S + (Q*/2).i.P", italic=True)
    add_paragraph(doc, "Bước 4: So sánh các TC vừa tính. Cái nào RẺ NHẤT (TC Min) => Kết luận: Mua số lượng Q đó ở mức giá P đó.")
    
    add_paragraph(doc, "⚠️ LƯU Ý CHẾT NGƯỜI (TRÁNH MẤT ĐIỂM):", bold=True, color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- Không được quên CỘNG TIỀN MUA HÀNG (D.P) vào công thức TC. Chỉ tính phí đặt + phí tồn trữ là SAI BẢN CHẤT bài toán chiết khấu.", color=RGBColor(255, 0, 0))

    add_heading(doc, "CÂU 7: QUẢN LÝ DỰ ÁN - MẠNG AON (CPM)", level=1)
    
    add_paragraph(doc, "1. Dịch Đề:", bold=True)
    add_paragraph(doc, "Đề cho cột Công việc (A, B, C...), cột CV Trực tiếp trước (Để vẽ đường), cột Thời gian t.")

    add_paragraph(doc, "2. Suy luận giải:", bold=True)
    add_paragraph(doc, "- Bắt đầu vẽ sơ đồ mạng AON (Các node là hình chữ nhật hoặc vòng tròn). Mũi tên từ trái sang phải.")
    add_paragraph(doc, "- Tính ES (Sớm nhất bắt đầu) và EF (Sớm nhất kết thúc): ĐI TỚI.", bold=True)
    add_paragraph(doc, "  => Quy tắc: EF = ES + t. Nếu 1 công việc nhận 2 mũi tên tới, ES = MAX(các EF trước đó).")
    add_paragraph(doc, "- Tính LF (Muộn nhất kết thúc) và LS (Muộn nhất bắt đầu): ĐI LÙI.", bold=True)
    add_paragraph(doc, "  => Quy tắc: LS = LF - t. Bắt đầu từ công việc cuối cùng (LF = EF của DA). Đi lùi về, nếu 1 nút trả 2 mũi tên về trước, LF = MIN(các LS sau đó).")
    add_paragraph(doc, "- Tính thời gian dự trữ (S hoặc Slack):", bold=True)
    add_paragraph(doc, "  => S = LS - ES (hoặc LF - EF).")
    add_paragraph(doc, "- Đường găng (Critical Path):", bold=True)
    add_paragraph(doc, "  => Những CV có S = 0 chính là đường găng. (Thời gian hoàn thành DA là độ dài đường này).")
    
    add_paragraph(doc, "- Vẽ Biểu đồ Gantt (Phương án sớm):", bold=True)
    add_paragraph(doc, "  => Vẽ trục ngang là Số ngày (0, 1, 2, 3...), trục dọc là Tên Công việc (A, B, C...).")
    add_paragraph(doc, "  => Tìm giá trị ES (Ngày bắt đầu sớm nhất) của công việc, bắt đầu vẽ thanh ngang từ ngày đó, kéo dài bằng đúng số ngày thực hiện (t) của công việc.")

    add_paragraph(doc, "⚠️ LƯU Ý CHẾT NGƯỜI (TRÁNH MẤT ĐIỂM):", bold=True, color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- Đi TỚI chọn MAX, đi LÙI chọn MIN. Sai một nút ở đầu là sai toàn bộ bảng, tính thật kỹ.", color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- Câu hỏi phụ \"Nếu CV X chậm Y ngày...\": So sánh Y với số ngày dự trữ S của CV X. Y ≤ S thì DA không đổi, Y > S thì DA bị trễ.", color=RGBColor(255, 0, 0))

    add_heading(doc, "CÂU 8: QUẢN LÝ DỰ ÁN - PERT & XÁC SUẤT", level=1)
    
    add_paragraph(doc, "1. Dịch Đề:", bold=True)
    add_paragraph(doc, "Đề không cho Thời gian cụ thể mà cho 3 loại thời gian: a (lạc quan), m (thường gặp), b (bi quan). Yêu cầu tính xác suất.")

    add_paragraph(doc, "2. Suy luận giải:", bold=True)
    add_paragraph(doc, "Bước 1: Tính lại Thời gian kỳ vọng (t) và Phương sai (V) cho TẤT CẢ công việc.")
    add_paragraph(doc, "  => t = (a + 4m + b) / 6")
    add_paragraph(doc, "  => V = ((b - a) / 6)²")
    add_paragraph(doc, "Bước 2: Vẽ mạng AON dựa vào (t) vừa tính, tìm ra Đường găng và thời gian hoàn thành dự án (T_E).")
    add_paragraph(doc, "Bước 3: CÓ ĐƯỜNG GĂNG RỒI mới được tính Phương sai Dự án (V_p).")
    add_paragraph(doc, "  => Cộng các V của những công việc NẰM TRÊN ĐƯỜNG GĂNG lại. (Tuyệt đối không cộng CV ngoài đường găng).")
    add_paragraph(doc, "Bước 4: Tính Độ lệch chuẩn σ = √V_p")
    add_paragraph(doc, "Bước 5: Tính xác suất hoàn thành DA trong T ngày.")
    add_paragraph(doc, "  => Áp dụng công thức Z = (T - T_E) / σ")
    add_paragraph(doc, "  => Tra bảng chuẩn tắc ra số %.")

    add_paragraph(doc, "⚠️ LƯU Ý CHẾT NGƯỜI (TRÁNH MẤT ĐIỂM):", bold=True, color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- PHƯƠNG SAI DỰ ÁN (V_p): Nhắc lại lần nữa, chỉ cộng phương sai của các công việc TRÊN ĐƯỜNG GĂNG. Rất nhiều bạn cộng tổng hết cột V là sai bét.", color=RGBColor(255, 0, 0))
    add_paragraph(doc, "- Muốn đạt X% thì cần bao nhiêu ngày: Suy ngược. Tra bảng từ số X% đó tìm Z, sau đó thay vào T = T_E + Z.σ là ra số ngày.", color=RGBColor(255, 0, 0))

    # Save document
    out_path = os.path.join(os.getcwd(), 'docs', 'QTKDCKS', 'Giai_chi_tiet_De_cuong_QTKDCKS_SuyLuan.docx')
    doc.save(out_path)
    print(f"File saved successfully to {out_path}")

if __name__ == '__main__':
    main()