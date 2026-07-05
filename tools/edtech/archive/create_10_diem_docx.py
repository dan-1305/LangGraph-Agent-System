import sys
from docx import Document
from docx.shared import Pt, RGBColor

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('BÀI GIẢI MẪU - ĐẠT ĐIỂM TỐI ĐA (10/10) - MÔN BMTT', 0)
    doc.add_paragraph('Tài liệu trình bày các bước giải chuẩn sư phạm: Ngắn gọn, chặt chẽ, không bị trừ điểm trình bày.')

    # CÂU 1
    doc.add_heading('Câu 1: Sử dụng Mã Affine giải mã bản mã "rmpjvpwtwtmhgmpwtgfwgkzn" với a=19, b=23', level=1)
    doc.add_paragraph('1. Công thức giải mã Affine: P = a^(-1) * (C - b) mod 26.')
    doc.add_paragraph('2. Tìm phần tử nghịch đảo a^(-1) của 19 theo modulo 26:')
    doc.add_paragraph('Ta cần tìm x sao cho 19 * x ≡ 1 mod 26.')
    doc.add_paragraph('Bằng thuật toán Euclide mở rộng (hoặc thử chọn), ta có: 19 * 11 = 209 = 8 * 26 + 1. Do đó a^(-1) = 11.')
    doc.add_paragraph('3. Thay vào công thức: P = 11 * (C - 23) mod 26.')
    doc.add_paragraph('Biến đổi (C - 23) ≡ (C + 3) mod 26. Suy ra công thức tính toán: P = 11 * (C + 3) mod 26.')
    doc.add_paragraph('4. Giải mã từng ký tự (A=0, B=1...):')
    
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Ký tự (C)'
    hdr[1].text = 'Số (C)'
    hdr[2].text = 'C + 3'
    hdr[3].text = '11 * (C + 3)'
    hdr[4].text = 'P (mod 26)'
    hdr[5].text = 'Bản rõ (P)'
    
    data = [('r',17,20,220,12,'M'), ('m',12,15,165,9,'J'), ('p',15,18,198,16,'Q'), ('j',9,12,132,2,'C'), ('v',21,24,264,4,'E')]
    for i, d in enumerate(data):
        cells = table.rows[i+1].cells
        for j in range(6):
            cells[j].text = str(d[j])
            
    doc.add_paragraph('... (Tính tương tự cho các ký tự còn lại)')
    doc.add_paragraph('=> Bản rõ: MJQCEQPIPIJGVJQPIVKPVNWU')

    # CÂU 2
    doc.add_heading('Câu 2: Mã hóa & giải mã Hill "doikhongnhulamo" với khóa K = [[7, 3], [8, 7]]', level=1)
    doc.add_paragraph('A. MÃ HÓA')
    doc.add_paragraph('1. Bản rõ dài 15 ký tự, lẻ nên thêm "x" vào cuối thành 16 ký tự. Chia làm 8 cặp vector cột:')
    doc.add_paragraph('do=(3,14), ik=(8,10), ho=(7,14), ng=(13,6), nh=(13,7), ul=(20,11), am=(0,12), ox=(14,23).')
    doc.add_paragraph('2. Áp dụng công thức C = K * P mod 26 (P là vector cột).')
    doc.add_paragraph('- Block 1 (do): [[7, 3], [8, 7]] * [3, 14]^T = [63, 122]^T ≡ [11, 18]^T mod 26 => ls')
    doc.add_paragraph('- Tính tương tự cho các block tiếp theo (Kết quả số -> chữ):')
    doc.add_paragraph('  + Block 2 (ik): [8, 4]^T => ie')
    doc.add_paragraph('  + Block 3 (ho): [13, 24]^T => ny')
    doc.add_paragraph('  + Block 4 (ng): [5, 16]^T => fq')
    doc.add_paragraph('  + Block 5 (nh): [8, 23]^T => ix')
    doc.add_paragraph('  + Block 6 (ul): [17, 3]^T => rd')
    doc.add_paragraph('  + Block 7 (am): [10, 6]^T => kg')
    doc.add_paragraph('  + Block 8 (ox): [11, 13]^T => ln')
    doc.add_paragraph('=> Bản mã tổng: lsienyfqixrdkgln')

    doc.add_paragraph('B. GIẢI MÃ')
    doc.add_paragraph('1. Tính ma trận K^(-1) = det(K)^(-1) * Adj(K) mod 26.')
    doc.add_paragraph('- det(K) = 7*7 - 3*8 = 25 ≡ -1 mod 26.')
    doc.add_paragraph('- Nghịch đảo det(K)^(-1): Vì (-1)*(-1) = 1 nên det(K)^(-1) = -1 ≡ 25 mod 26.')
    doc.add_paragraph('- Adj(K) = [[7, -3], [-8, 7]] ≡ [[7, 23], [18, 7]] mod 26.')
    doc.add_paragraph('=> K^(-1) = 25 * [[7, 23], [18, 7]] ≡ [[19, 3], [8, 19]] mod 26.')
    doc.add_paragraph('2. Áp dụng công thức P = K^(-1) * C mod 26.')
    doc.add_paragraph('- Block 1 (ls): [[19, 3], [8, 19]] * [11, 18]^T = [263, 430]^T ≡ [3, 14]^T mod 26 => do')
    doc.add_paragraph('- (Tính tương tự cho các block tiếp theo)')
    doc.add_paragraph('=> Bản rõ tổng: doikhongnhulamo(x)')

    # CÂU 3
    doc.add_heading('Câu 3: Sử dụng Mã Hill mã hóa "khoacongnghekythuat" với K = [[5, 8, 9], [3, 7, 5], [6, 3, 9]]', level=1)
    doc.add_paragraph('A. MÃ HÓA')
    doc.add_paragraph('1. Bản rõ: Thêm 2 "x" thành 21 ký tự, chia 7 vector 3x1: kho=(10,7,14), aco=(0,2,14), ngn=(13,6,13), ghe=(6,7,4), kyt=(10,24,19), hua=(7,20,0), txx=(19,23,23).')
    doc.add_paragraph('2. Tính C = K * P mod 26.')
    doc.add_paragraph('- Block 1 (kho): K * [10, 7, 14]^T = [232, 149, 207]^T ≡ [24, 19, 25]^T mod 26 => ytz')
    doc.add_paragraph('- Các block tiếp theo:')
    doc.add_paragraph('  + Block 2 (aco): [12, 6, 2]^T => mgc')
    doc.add_paragraph('  + Block 3 (ngn): [22, 16, 5]^T => wqf')
    doc.add_paragraph('  + Block 4 (ghe): [18, 9, 15]^T => sjp')
    doc.add_paragraph('  + Block 5 (kyt): [23, 7, 17]^T => xhr')
    doc.add_paragraph('  + Block 6 (hua): [13, 5, 24]^T => nfy')
    doc.add_paragraph('  + Block 7 (txx): [18, 21, 0]^T => sva')
    doc.add_paragraph('=> Bản mã tổng: ytzmgcwqfsjpxhrnfysva')

    doc.add_paragraph('B. GIẢI MÃ')
    doc.add_paragraph('1. Tính định thức det(K) (Có thể bấm máy tính trực tiếp):')
    doc.add_paragraph('det(K) = -33 ≡ (-33 + 52) = 19 mod 26.')
    doc.add_paragraph('2. Tìm det(K)^(-1): Ta có 19 * 11 = 209 ≡ 1 mod 26 => det(K)^(-1) = 11.')
    doc.add_paragraph('3. Tính ma trận phụ hợp Adj(K) bằng phần bù đại số A_ij = (-1)^(i+j) * det(M_ij):')
    doc.add_paragraph('- A_11 = (-1)^(1+1) * (7*9 - 5*3) = 48 ≡ 22 mod 26.')
    doc.add_paragraph('- A_12 = (-1)^(1+2) * (3*9 - 5*6) = -(-3) = 3 ≡ 3 mod 26.')
    doc.add_paragraph('- (Tính tương tự cho các phần tử còn lại, sau đó chuyển vị ma trận A_ij)')
    doc.add_paragraph('=> Adj(K) = [[22, 9, 3], [3, 17, 2], [5, 1, 11]]')
    doc.add_paragraph('4. K^(-1) = 11 * Adj(K) mod 26 = [[8, 25, 7], [7, 5, 22], [1, 25, 17]].')
    doc.add_paragraph('5. P = K^(-1) * C mod 26. Block 1: K^(-1) * [24, 19, 25]^T = [842, 813, 924]^T ≡ [10, 7, 14]^T mod 26 => kho.')

    # CÂU 4
    doc.add_heading('Câu 4: Mã hóa RSA nội dung "chuccacbanthitotnhe" với p=11, q=13', level=1)
    doc.add_paragraph('Bước 1: Tính n = p * q = 143. Số Euler φ(n) = (p-1)*(q-1) = 120.')
    doc.add_paragraph('Bước 2: Chọn e = 7 (thỏa mãn gcd(7, 120) = 1). Khóa công khai PU=(7, 143).')
    doc.add_paragraph('Bước 3: Tìm khóa bí mật d thỏa mãn d * e ≡ 1 mod 120 (Tức là 7 * d ≡ 1 mod 120).')
    doc.add_paragraph('Bằng thuật toán Euclide mở rộng (hoặc thử chọn), ta tìm được d = 103 vì 7 * 103 = 721 = 120 * 6 + 1. Khóa bí mật PR=(103, 143).')
    doc.add_paragraph('Bước 4: Mã hóa bản rõ C = M^e mod 143.')
    doc.add_paragraph('- Ký tự "c" (M=2): C1 = 2^7 mod 143 = 128 mod 143 = 128.')
    doc.add_paragraph('- Ký tự "h" (M=7): C2 = 7^7 mod 143 = 823543 ≡ 6 mod 143.')
    doc.add_paragraph('- Ký tự "u" (M=20): C3 = 20^7 mod 143 ≡ 136 mod 143.')
    doc.add_paragraph('=> Mã hóa tương tự: [128, 6, 136, 128, 128, 0, 128, 1...]')

    # CÂU 6 & 7 (PLAYFAIR)
    doc.add_heading('Câu 6 & 7: Mã hóa Playfair', level=1)
    doc.add_paragraph('Bước 1: Lập ma trận 5x5 từ từ khóa (Bỏ ký tự trùng, gộp I và J).')
    doc.add_paragraph('- Với key "bainaylamroi" -> B A I N Y | L M R O C | D E F G H | K P Q S T | U V W X Z.')
    doc.add_paragraph('- Với key "lamnhieuroi" -> L A M N H | I E U R O | B C D F G | K P Q S T | V W X Y Z.')
    doc.add_paragraph('Bước 2: Tách cặp (Ví dụ tên NGUYEN VAN A -> NG UY EN VA NA).')
    doc.add_paragraph('Bước 3: Áp dụng quy tắc (Khác hàng/cột lấy góc HCN; Cùng hàng dịch phải; Cùng cột dịch xuống).')
    doc.add_paragraph('Ví dụ cặp "NG" (Key "bainaylamroi"): N(Hàng 1) và G(Hàng 3) -> Lấy góc hình chữ nhật đối diện -> YF.')

    # CÁC CÂU LÝ THUYẾT
    doc.add_heading('Câu 5, 8, 9, 10: Lý thuyết', level=1)
    doc.add_paragraph('Câu 5 (IPS): Là hệ thống phòng chống xâm nhập CHỦ ĐỘNG. Nằm trực tiếp trên luồng dữ liệu (Inline). Có khả năng phát hiện hành vi độc hại và ngắt kết nối/Drop gói tin ngay lập tức.')
    doc.add_paragraph('Câu 8 (IDS vs IPS): IDS (Thụ động, đứng ngoài quan sát, chỉ cảnh báo/Alert, không tự động chặn). IPS (Chủ động, ngăn chặn trực tiếp các cuộc tấn công).')
    doc.add_paragraph('Câu 9 (SSL/TLS): Cung cấp 3 tính năng cốt lõi: 1) Encryption (Mã hóa đường truyền). 2) Authentication (Xác thực máy chủ qua Chứng chỉ số). 3) Integrity (Toàn vẹn dữ liệu).')
    doc.add_paragraph('Câu 10 (Chữ ký số): Người gửi dùng hàm Hash băm tài liệu -> Mã hóa bản Hash bằng Khóa bí mật (Private Key) để tạo chữ ký. Người nhận dùng Khóa công khai (Public Key) giải mã -> Băm tài liệu nhận được và so sánh 2 bản Hash để xác thực. Tác dụng: Chống giả mạo, toàn vẹn dữ liệu, chống chối bỏ.')

    doc.save('docs/BMTT/Giai_Thi_10_Diem.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()