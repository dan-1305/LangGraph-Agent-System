import sys
from docx import Document
from docx.shared import Pt, RGBColor

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('BẢNG CÔNG THỨC CỐT LÕI (CHEATSHEET) - QTKDCKS', 0)
    doc.add_paragraph('Mang bảng này vào phòng thi để ốp công thức cho 14 bài tập thực hành.')

    # 1. QUYẾT ĐỊNH RỦI RO
    doc.add_heading('1. RA QUYẾT ĐỊNH (CHƯƠNG 2)', level=1)
    doc.add_paragraph('- EMV (Giá trị tiền tệ kỳ vọng) = Lợi nhuận * Xác suất + Lợi nhuận * Xác suất... -> CHỌN MAX EMV.')
    doc.add_paragraph('- EOL (Tổn thất cơ hội kỳ vọng):')
    doc.add_paragraph('  + Bước 1: Lập ma trận tiếc nuối (Lấy Max cột - Các giá trị trong cột đó).')
    doc.add_paragraph('  + Bước 2: EOL = Giá trị tiếc nuối * Xác suất... -> CHỌN MIN EOL.')
    doc.add_paragraph('- EVPI (Giá trị thông tin hoàn hảo) = Min EOL.')
    doc.add_paragraph('- Maximax: Lạc quan -> CHỌN MAX CỦA NHỮNG CÁI MAX.')
    doc.add_paragraph('- Maximin: Bi quan -> CHỌN MAX CỦA NHỮNG CÁI MIN (Chọn lỗ ít nhất).')
    doc.add_paragraph('- Hurwicz (hệ số α): H = α * Max + (1 - α) * Min -> CHỌN MAX H.')

    # 2. TÀI CHÍNH & KẾ TOÁN
    doc.add_heading('2. KẾ TOÁN VÀ TÀI CHÍNH (CHƯƠNG 3, 4)', level=1)
    doc.add_paragraph('- TỔNG TÀI SẢN = TỔNG NỢ + VỐN CHỦ SỞ HỮU.')
    doc.add_paragraph('- Lãi gộp = Doanh thu - Giá vốn hàng bán.')
    doc.add_paragraph('- LNTT (Lợi nhuận trước thuế) = Lãi gộp - Chi phí HĐ - Lãi vay.')
    doc.add_paragraph('- LNST (Lợi nhuận sau thuế) = LNTT - Thuế (Hoặc LNTT * (1 - Thuế suất)).')
    doc.add_paragraph('- Lợi nhuận giữ lại cuối kỳ = LNGL đầu kỳ + LNST - Cổ tức.')
    doc.add_paragraph('- Tỷ số nợ = Tổng nợ / Tổng tài sản.')
    doc.add_paragraph('- Thanh toán hiện hành (Current ratio) = Tài sản lưu động / Nợ ngắn hạn.')
    doc.add_paragraph('- Thanh toán nhanh (Quick ratio) = (Tài sản lưu động - Hàng tồn kho) / Nợ ngắn hạn.')
    doc.add_paragraph('- Vòng quay tài sản = Doanh thu / Tổng tài sản.')
    doc.add_paragraph('- Kỳ thu tiền bình quân = Khoản phải thu / (Doanh thu / 365). Hoặc lấy Khoản phải thu / Doanh thu bình quân ngày.')

    # 3. QUẢN LÝ HÀNG TỒN KHO
    doc.add_heading('3. QUẢN LÝ HÀNG TỒN KHO (EOQ & ROP)', level=1)
    doc.add_paragraph('- Điểm đặt hàng tối ưu (EOQ) = √(2*D*S / H)')
    doc.add_paragraph('  (Trong đó: D = Nhu cầu năm, S = Chi phí 1 lần đặt hàng, H = Chi phí lưu trữ 1 đơn vị/năm).')
    doc.add_paragraph('- Số lần đặt hàng (N) = D / EOQ.')
    doc.add_paragraph('- Chu kỳ đặt hàng (T) = Số ngày làm việc trong năm / N.')
    doc.add_paragraph('- Tổng chi phí (TC) = (D/EOQ)*S + (EOQ/2)*H.')
    doc.add_paragraph('- Nhu cầu bình quân ngày (d) = Nhu cầu năm / Số ngày làm việc.')
    doc.add_paragraph('- Điểm đặt hàng lại (ROP) = (d * Lead time) + Dự trữ an toàn (Nếu có).')

    # 4. QUẢN TRỊ DỰ ÁN
    doc.add_heading('4. QUẢN TRỊ DỰ ÁN (PERT/CPM - CHƯƠNG 7)', level=1)
    doc.add_paragraph('- Thời gian hoàn thành dự án = Tổng thời gian của NHÁNH DÀI NHẤT (Đường găng).')
    doc.add_paragraph('- Đường găng (Critical Path): Là nhánh dài nhất. Các công việc trên đường găng có Thời gian dự trữ (Slack) = 0.')
    doc.add_paragraph('- Thời gian dự trữ (Slack) = Thời gian hoàn thành dự án (Max) - Thời gian của nhánh chứa công việc đó.')
    doc.add_paragraph('- Thời gian kỳ vọng (Te) = (a + 4m + b) / 6.')
    doc.add_paragraph('- Phương sai (V) = ((b - a) / 6)^2.')

    doc.save('docs/QTKDCKS/Cheatsheet_QTKDCKS.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()