import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def init_document(title: str) -> Document:
    """Khởi tạo file Word, thiết lập tiêu đề chính."""
    doc = Document()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return doc

def add_affine_section(doc: Document) -> None:
    """Module tạo nội dung rút gọn cho Câu 1: Mã Affine."""
    doc.add_heading('Câu 1: Mã hóa Affine (a=19, b=23)', level=1)
    
    content = (
        "1. Tìm nghịch đảo a^(-1):\n"
        "Ta có: 19 * 11 = 209 = 8 * 26 + 1 => 19 * 11 ≡ 1 (mod 26). Vậy a^(-1) = 11.\n\n"
        "2. Công thức giải mã tối ưu:\n"
        "P = 11 * (C - 23) mod 26 ≡ 11 * (C + 3) mod 26.\n\n"
        "3. Kết quả giải mã (Ghi bảng nháp thẳng vào bài):\n"
        "Bản mã C : 17(r)  12(m)  15(p)  9(j)  21(v)\n"
        "Bản rõ P : 12(M)   9(J)  16(Q)  2(C)   4(E)\n"
        "=> Bản rõ tổng: MJQCEQPIPIJGVJQPIVKPVNWU"
    )
    doc.add_paragraph(content)

def add_hill_section(doc: Document) -> None:
    """Module tạo nội dung rút gọn cho Câu 3: Mã Hill 3x3."""
    doc.add_heading('Câu 3: Mã Hill 3x3 (Sinh tồn phòng thi)', level=1)
    
    content = (
        "1. Tính định thức det(K):\n"
        "det(K) = -33. Do không được để số âm: -33 + 26*2 = 19.\n\n"
        "2. Tìm nghịch đảo det(K)^(-1):\n"
        "Ta có 19 * 11 = 209 = 8 * 26 + 1. Vậy det(K)^(-1) = 11.\n\n"
        "3. Ma trận phụ hợp Adj(K) (Chỉ ghi 1 ví dụ, còn lại phang kết quả):\n"
        "C_11 = (-1)^(1+1) * (7*9 - 5*3) = 48 ≡ 22 (mod 26).\n"
        "=> Adj(K) = [[22, 19, 3], [3, 9, 2], [19, 1, 11]] (mod 26).\n\n"
        "4. K^(-1) = 11 * Adj(K) mod 26."
    )
    doc.add_paragraph(content)

def add_playfair_rsa_theory(doc: Document) -> None:
    """Module tạo nội dung rút gọn Playfair, RSA và Lý thuyết."""
    # Playfair
    doc.add_heading('Câu 6 & 7: Mã Playfair', level=1)
    doc.add_paragraph(
        "Trình bày 3 bước ngắn gọn:\n"
        "1. Vẽ ma trận khóa 5x5 (Gộp I/J).\n"
        "2. Cặp bản rõ: NG UY EN VA NA\n"
        "3. Cặp bản mã: YF ZB HA AV HM (Ghi thẳng kết quả theo quy tắc, không diễn giải dài dòng chữ nghĩa)."
    )
    
    # RSA
    doc.add_heading('Câu 4: Mã hóa RSA (p=11, q=13)', level=1)
    doc.add_paragraph(
        "n = 143; Φ(n) = 120.\n"
        "Chọn e = 7 (thỏa gcd(7, 120) = 1).\n"
        "Khóa d: 7 * 103 = 721 = 120 * 6 + 1 => d = 103.\n"
        "Lưu ý: Dùng thuật toán bình phương và nhân để tính mũ lớn (VD: 7^7 = (7^3)^2 * 7)."
    )
    
    # Theory
    doc.add_heading('Lý thuyết cốt lõi (Câu 5, 8, 9, 10)', level=1)
    doc.add_paragraph(
        "- IPS vs IDS: IPS chặn chủ động (Inline), IDS chỉ cảnh báo thụ động.\n"
        "- SSL/TLS: Encryption (Mã hóa), Authentication (Xác thực), Integrity (Toàn vẹn).\n"
        "- Chữ ký số: Băm tài liệu (Hash) -> Mã hóa Hash bằng Private Key của người gửi."
    )

def generate_exam_doc(filename: str) -> None:
    """Hàm tổng phối hợp các module để ráp thành file hoàn chỉnh."""
    doc = init_document("TÀI LIỆU ÔN THI BMTT - BẢN RÚT GỌN (TỐI ƯU THỜI GIAN)")
    
    add_affine_section(doc)
    add_hill_section(doc)
    add_playfair_rsa_theory(doc)
    
    doc.save(filename)
    print(f"[+] Đã tạo thành công file: {filename}")

# ==========================================
# QUALITY GATE: UNIT TEST
# ==========================================
def run_unit_test() -> None:
    """Test nhanh xem file có thực sự được tạo ra và lưu thành công không."""
    test_file = "test_bmtt_output.docx"
    
    print("[*] Chạy Unit Test...")
    generate_exam_doc(test_file)
    
    assert os.path.exists(test_file), "Lỗi: File Word chưa được tạo!"
    print(f"[+] Unit Test Passed: File {test_file} đã được tạo.")
    
    # Dọn dẹp file test
    os.remove(test_file)
    print("[*] Đã dọn dẹp file test.\n")

if __name__ == "__main__":
    # 1. Chạy Quality Gate trước
    run_unit_test()
    
    # 2. Sinh file thật
    final_filename = "BMTT_Bi_Kip_Phong_Thi.docx"
    generate_exam_doc(final_filename)