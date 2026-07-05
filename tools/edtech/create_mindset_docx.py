from docx import Document
from docx.shared import Pt, RGBColor

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('MINDSET & CÁCH TRÌNH BÀY GIẤY THI BMTT', 0)
    doc.add_paragraph('Tài liệu này mô phỏng luồng suy nghĩ (Mindset) khi ngồi trong phòng thi và cách viết ra giấy sao cho ngắn gọn, đủ điểm.')

    # CÂU 1
    doc.add_heading('CÂU 1: GIẢI MÃ AFFINE', level=1)
    
    p_mindset = doc.add_paragraph()
    run = p_mindset.add_run('💭 Trong đầu nghĩ (Nháp):')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192) # Blue
    doc.add_paragraph('- Đề kêu giải mã Affine a=19, b=23.')
    doc.add_paragraph('- Công thức giải mã là ngược của mã hóa: P = a^(-1) * (C - b) mod 26.')
    doc.add_paragraph('- Tìm a^(-1): Bấm máy tính 19*11 = 209. Lấy 209 chia 26 dư 1. Vậy a^(-1) = 11.')
    doc.add_paragraph('- Tính (C - 23): Dễ sai số âm, nên đổi -23 thành +3 (vì 26 - 23 = 3). Vậy (C - 23) ≡ (C + 3) mod 26.')
    doc.add_paragraph('- Chốt công thức bấm máy: P = 11 * (C + 3) mod 26. Giờ chỉ việc kẻ bảng ráp số vào bấm.')

    p_write = doc.add_paragraph()
    run = p_write.add_run('✍️ Viết ra giấy thi:')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80) # Green
    doc.add_paragraph('Bước 1: Công thức giải mã Affine: P = a^(-1) * (C - b) mod 26.')
    doc.add_paragraph('Bước 2: Tìm phần tử nghịch đảo của a=19. Ta có 19 * 11 ≡ 1 mod 26 => a^(-1) = 11.')
    doc.add_paragraph('Bước 3: Suy ra công thức: P = 11 * (C - 23) ≡ 11 * (C + 3) mod 26.')
    doc.add_paragraph('Bước 4: Kẻ bảng tính (C -> số -> P -> chữ). Kết luận bản rõ là: MJQCEQPIPIJGVJQPIVKVWNU.')

    # CÂU 2
    doc.add_heading('CÂU 2: MÃ HILL (2x2)', level=1)
    
    p_mindset = doc.add_paragraph()
    run = p_mindset.add_run('💭 Trong đầu nghĩ (Nháp):')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)
    doc.add_paragraph('- Đề mã hóa Hill ma trận 2x2. "doikhongnhulamo" có 15 chữ, lẻ rồi. Thêm chữ "x" vào cuối cho chẵn 16 chữ (8 cặp).')
    doc.add_paragraph('- Mã hóa thì cứ lấy từng cặp chữ nhân với ma trận K. Nhớ là (dòng 1x2) x (ma trận 2x2) theo modulo 26.')
    doc.add_paragraph('- Giải mã mới chua: Phải tính định thức det(K) = 7*7 - 3*8 = 25.')
    doc.add_paragraph('- Nghịch đảo của 25 theo mod 26 là số mấy? 25 = -1 mod 26. Mà (-1)*(-1)=1. Vậy nghịch đảo là 25.')
    doc.add_paragraph('- Tìm K^(-1): Đổi chéo chính (7,7 giữ nguyên), đổi dấu chéo phụ (-3, -8). Rồi đem nhân với 25.')
    doc.add_paragraph('- Có K^(-1) rồi thì lấy bản mã nhân K^(-1) là ra lại bản rõ.')

    p_write = doc.add_paragraph()
    run = p_write.add_run('✍️ Viết ra giấy thi:')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80)
    doc.add_paragraph('A. MÃ HÓA')
    doc.add_paragraph('Bước 1: Thêm "x" vào cuối bản rõ. Tách thành các block 2 ký tự: (do), (ik)...')
    doc.add_paragraph('Bước 2: C = P x K mod 26. Block 1: [3, 14] x [[7,3], [8,7]] = [133, 107] ≡ [3, 3] mod 26 => DD. (Làm tương tự ra bản mã).')
    doc.add_paragraph('B. GIẢI MÃ')
    doc.add_paragraph('Bước 1: det(K) = 7*7 - 3*8 = 25 ≡ -1 mod 26.')
    doc.add_paragraph('Bước 2: det(K)^(-1) ≡ -1 ≡ 25 mod 26.')
    doc.add_paragraph('Bước 3: Ma trận phụ hợp Adj(K) = [[7, -3], [-8, 7]] ≡ [[7, 23], [18, 7]] mod 26.')
    doc.add_paragraph('Bước 4: K^(-1) = 25 * Adj(K) = [[19, 3], [8, 19]] mod 26.')
    doc.add_paragraph('Bước 5: P = C x K^(-1) mod 26. Block 1: [3,3] x [[19,3], [8,19]] = [81, 66] ≡ [3, 14] mod 26 => do.')

    # CÂU 4
    doc.add_heading('CÂU 4: MÃ HÓA RSA', level=1)
    
    p_mindset = doc.add_paragraph()
    run = p_mindset.add_run('💭 Trong đầu nghĩ (Nháp):')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)
    doc.add_paragraph('- Mã hóa RSA với p=11, q=13. Mình cần tìm n, φ(n), chọn e và tính d.')
    doc.add_paragraph('- Tính n = 11*13=143. Tính φ(n) = 10*12=120.')
    doc.add_paragraph('- Mình được tự chọn e, chọn đại e=7 cho dễ tính. Khóa công khai là (7, 143).')
    doc.add_paragraph('- Tính d sao cho 7*d = 1 mod 120. Bấm nhẩm: 7*103 = 721 = 120*6 + 1 => d=103.')
    doc.add_paragraph('- Mã hóa chữ "c" (số 2): Tính 2^7 mod 143. 2^7 = 128 < 143 nên C = 128.')

    p_write = doc.add_paragraph()
    run = p_write.add_run('✍️ Viết ra giấy thi:')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80)
    doc.add_paragraph('Bước 1: n = p*q = 143. φ(n) = (p-1)(q-1) = 120.')
    doc.add_paragraph('Bước 2: Chọn e = 7 thỏa mãn gcd(7, 120) = 1.')
    doc.add_paragraph('Bước 3: Tính d ≡ 7^(-1) mod 120. Chọn d = 103 (vì 7*103 ≡ 1 mod 120).')
    doc.add_paragraph('Bước 4: Mã hóa C = M^e mod 143. Với chữ c (2) => C = 2^7 mod 143 = 128. (Viết tương tự 1 vài ký tự đại diện là đủ).')

    # CÂU 6
    doc.add_heading('CÂU 6: PLAYFAIR', level=1)
    
    p_mindset = doc.add_paragraph()
    run = p_mindset.add_run('💭 Trong đầu nghĩ (Nháp):')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)
    doc.add_paragraph('- Khóa "bainaylamroi", bỏ chữ trùng: b a i n y l m r o. Lập bảng 5x5 điền nốt chữ cái.')
    doc.add_paragraph('- Tên LE VAN AN (giả sử). Chia cặp LE VA NA Nx.')
    doc.add_paragraph('- Quy tắc: Cùng hàng thì tiến 1 ô sang phải. Cùng cột thì tiến 1 ô xuống. Khác hàng/cột thì lấy góc hình chữ nhật đối diện.')

    p_write = doc.add_paragraph()
    run = p_write.add_run('✍️ Viết ra giấy thi:')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80)
    doc.add_paragraph('Bước 1: Lập ma trận 5x5 từ khóa.')
    doc.add_paragraph('Bước 2: Tách tên thành các block 2 ký tự (có chêm X nếu cần).')
    doc.add_paragraph('Bước 3: Ghi rõ quy tắc dò bảng (cùng hàng/cột/hình chữ nhật) và suy ra bản mã.')

    doc.save('docs/BMTT/Mindset_Lam_Bai_BMTT.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()
