import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = None

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('BẢN RÚT GỌN ÔN TẬP QTKD (CHEAT SHEET)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True

    add_heading(doc, "PHẦN 1: LÝ THUYẾT (TỪ KHÓA ĂN ĐIỂM)", level=1)

    # Câu 1
    add_heading(doc, "Câu 1: Các loại ra quyết định", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Chắc chắn: ").bold = True
    p.add_run("Biết trước 100% kết quả. => Chọn PA tốt nhất.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Rủi ro: ").bold = True
    p.add_run("Biết được xác suất xảy ra của từng kết quả. => Dùng Giá trị kỳ vọng (EMV).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Không chắc chắn: ").bold = True
    p.add_run("Không biết kết quả và không biết xác suất. => Phụ thuộc thái độ người ra quyết định.")

    # Câu 2
    add_heading(doc, "Câu 2: Các mô hình quyết định không chắc chắn", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Maximax (Lạc quan): ").bold = True
    p.add_run("Tìm Max của từng PA -> Chọn Max lớn nhất (Max của Max).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Maximin (Bi quan): ").bold = True
    p.add_run("Tìm Min của từng PA -> Chọn Min lớn nhất (Max của Min).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Hurwicz (Thực tế): ").bold = True
    p.add_run("Dùng hệ số lạc quan α. Công thức: H = α.Max + (1-α).Min -> Chọn H lớn nhất.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Minimax (Tiếc nuối): ").bold = True
    p.add_run("Lập bảng tiếc nuối -> Tìm mức tiếc nuối Max của từng PA -> Chọn Max nhỏ nhất (Min của Max).")

    # Câu 3
    add_heading(doc, "Câu 3: Quản lý công nghệ và Mô hình THIO", level=2)
    add_paragraph(doc, "- Quản lý công nghệ: Kết nối quản lý, kỹ thuật và khoa học để khai thác năng lực công nghệ.")
    add_paragraph(doc, "- 4 Yếu tố cấu thành (THIO):", bold=True)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("T (Technoware - Kỹ thuật): ").bold = True
    p.add_run("Máy móc, công cụ, dây chuyền (Phần cứng).")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("H (Humanware - Con người): ").bold = True
    p.add_run("Kỹ năng, kinh nghiệm, năng lực sáng tạo của người lao động.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("I (Inforware - Thông tin): ").bold = True
    p.add_run("Dữ liệu, bản vẽ, quy trình, phần mềm (Phần mềm).")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("O (Orgaware - Tổ chức): ").bold = True
    p.add_run("Cơ cấu quản lý, mạng lưới phối hợp.")

    # Câu 4
    add_heading(doc, "Câu 4: Marketing Mix (4P)", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Product (Sản phẩm): ").bold = True
    p.add_run("Hàng hóa/dịch vụ thỏa mãn nhu cầu (chất lượng, thiết kế, bảo hành).")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Price (Giá cả): ").bold = True
    p.add_run("Chữ P tạo ra doanh thu (chiến lược giá, chiết khấu).")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Place (Phân phối): ").bold = True
    p.add_run("Đưa sản phẩm đến tay KH (Kênh trực tiếp/gián tiếp, vị trí).")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Promotion (Xúc tiến): ").bold = True
    p.add_run("Truyền thông (Quảng cáo, khuyến mãi, PR).")

    doc.add_page_break()

    # Phụ lục Công thức
    add_heading(doc, "PHẦN 2: TỔNG HỢP CÔNG THỨC (HỌC THUỘC)", level=1)

    add_heading(doc, "1. QUẢN TRỊ TỒN KHO", level=2)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("EOQ (Lượng đặt hàng kinh tế): ").bold = True
    p.add_run("EOQ = √(2.D.S / H)")
    p.add_run("\nTrong đó: D: Nhu cầu năm | S: Phí đặt hàng 1 lần | H: Phí tồn trữ 1 đơn vị/năm (Nếu cho tỷ lệ i% thì H = i.P, P là giá mua).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Số lần đặt hàng trong năm (N): ").bold = True
    p.add_run("N = D / EOQ")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Chu kỳ đặt hàng (T): ").bold = True
    p.add_run("T = Số ngày làm việc / N")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tổng chi phí (TC): ").bold = True
    p.add_run("\nTC = Phí Mua (D.P) + Phí Đặt hàng ((D/Q).S) + Phí Tồn trữ ((Q/2).H)")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Điểm tái đặt hàng (ROP): ").bold = True
    p.add_run("\nROP = d × L + SS")
    p.add_run("\nTrong đó: d: nhu cầu/ngày (D/Số ngày làm việc) | L: Thời gian giao hàng | SS: Dự trữ an toàn.")

    add_heading(doc, "2. QUẢN LÝ DỰ ÁN", level=2)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Thông số mạng (AON): ").bold = True
    p.add_run("\n- Chiều đi (ES, EF): ES = EF_max_trước | EF = ES + Thời gian(t)")
    p.add_run("\n- Chiều về (LS, LF): LF = LS_min_sau | LS = LF - t")
    p.add_run("\n- Dự trữ (Slack S): S = LS - ES (hoặc LF - EF).")
    p.add_run("\n- Đường găng: Là đường đi qua các CV có S = 0.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Mô hình PERT: ").bold = True
    p.add_run("\n- Thời gian kỳ vọng (t): t = (a + 4m + b) / 6")
    p.add_run("\n- Phương sai 1 CV (V): V = ((b - a) / 6)²")
    p.add_run("\n- Phương sai DA (V_p): Cộng các V của CV NẰM TRÊN ĐƯỜNG GĂNG.")
    p.add_run("\n- Độ lệch chuẩn (σ): σ = √V_p")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Xác suất hoàn thành DA: ").bold = True
    p.add_run("\n- Tính Z = (T - T_E) / σ")
    p.add_run("\n  (T: T/gian đề cho | T_E: T/gian hoàn thành DA)")
    p.add_run("\n- Tra bảng Z để tìm % xác suất.")
    p.add_run("\n- Tính ngược (Muốn đạt X% thì mất bao lâu): Tra bảng ngược ra Z, thay vào: T = T_E + Z.σ")

    # Save document
    out_path = os.path.join(os.getcwd(), 'docs', 'QTKDCKS', 'Giai_chi_tiet_De_cuong_QTKDCKS_RutGon.docx')
    doc.save(out_path)
    print(f"File saved successfully to {out_path}")

if __name__ == '__main__':
    main()