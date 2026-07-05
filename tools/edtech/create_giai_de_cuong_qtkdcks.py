import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = None

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('GIẢI CHI TIẾT ĐỀ CƯƠNG ÔN TẬP QTKD CHO KỸ SƯ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True

    add_heading(doc, "PHẦN 1: LÝ THUYẾT", level=1)

    # Câu 1
    add_heading(doc, "Câu 1: Phân biệt các loại ra quyết định (Trong điều kiện chắc chắn, Rủi ro, Không chắc chắn)", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Ra quyết định trong điều kiện chắc chắn: ").bold = True
    p.add_run("Là tình huống mà người ra quyết định biết chắc chắn kết quả của từng phương án sẽ xảy ra. Thông tin là đầy đủ và chính xác (Xác suất xảy ra là 100%). Người ra quyết định chỉ việc chọn phương án mang lại lợi ích cao nhất (hoặc chi phí thấp nhất).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Ra quyết định trong điều kiện rủi ro: ").bold = True
    p.add_run("Là tình huống mà người ra quyết định không biết chắc chắn kết quả nào sẽ xảy ra cho mỗi phương án, nhưng có thể ước lượng được xác suất xảy ra của từng trạng thái tự nhiên (kết quả). Thông thường dựa vào dữ liệu lịch sử hoặc kinh nghiệm để xác định xác suất. Phương pháp phổ biến là tính Giá trị kỳ vọng (EMV).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Ra quyết định trong điều kiện không chắc chắn: ").bold = True
    p.add_run("Là tình huống mà người ra quyết định hoàn toàn không biết kết quả nào sẽ xảy ra và cũng không thể xác định hoặc ước lượng được xác suất xảy ra của các trạng thái tự nhiên. Quyết định phụ thuộc nhiều vào thái độ của người ra quyết định (lạc quan, bi quan,...).")

    # Câu 2
    add_heading(doc, "Câu 2: Nêu và giải thích các mô hình ra quyết định trong điều kiện không chắc chắn: Maximax, Maximin, Hurwicz, Minimax.", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tiêu chuẩn Maximax (Lạc quan): ").bold = True
    p.add_run("Người ra quyết định kỳ vọng điều tốt nhất sẽ xảy ra. Phương pháp: Tìm giá trị cực đại của từng phương án, sau đó chọn phương án có giá trị cực đại lớn nhất (Max của các Max).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tiêu chuẩn Maximin (Bi quan - Wald): ").bold = True
    p.add_run("Người ra quyết định cho rằng điều tồi tệ nhất sẽ xảy ra. Phương pháp: Tìm giá trị cực tiểu của từng phương án (kết quả xấu nhất), sau đó chọn phương án có giá trị cực tiểu lớn nhất (Max của các Min), nhằm đảm bảo một mức lợi ích tối thiểu.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tiêu chuẩn Hurwicz (Chủ nghĩa hiện thực): ").bold = True
    p.add_run("Là sự kết hợp giữa lạc quan và bi quan thông qua hệ số lạc quan α (0 ≤ α ≤ 1). Phương pháp: Tính giá trị Hurwicz cho từng phương án = α × (Giá trị Max) + (1 - α) × (Giá trị Min). Chọn phương án có giá trị lớn nhất.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tiêu chuẩn Minimax (Tiêu chuẩn tiếc nuối - Savage): ").bold = True
    p.add_run("Người ra quyết định muốn giảm thiểu sự tiếc nuối khi chọn sai phương án. Phương pháp: Lập bảng tiếc nuối (lấy giá trị lớn nhất của cột trừ đi các giá trị trong cột). Tìm mức tiếc nuối lớn nhất của mỗi phương án, sau đó chọn phương án có mức tiếc nuối lớn nhất nhỏ nhất (Min của các Max).")

    # Câu 3
    add_heading(doc, "Câu 3: Trình bày khái niệm quản lý công nghệ và các yếu tố cấu thành công nghệ.", level=2)
    add_paragraph(doc, "- Khái niệm quản lý công nghệ:", bold=True)
    add_paragraph(doc, "Quản lý công nghệ là tập hợp các hoạt động quản lý (lập kế hoạch, tổ chức, lãnh đạo, kiểm tra) liên quan đến việc tạo ra, thu nhận, duy trì và khai thác năng lực công nghệ nhằm đạt được các mục tiêu chiến lược của tổ chức. Nó kết nối các lĩnh vực kỹ thuật, khoa học và quản trị học.")
    add_paragraph(doc, "- Các yếu tố cấu thành công nghệ (Mô hình THIO):", bold=True)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Phần kỹ thuật (Technoware - T): ").bold = True
    p.add_run("Máy móc, thiết bị, công cụ, dây chuyền sản xuất, cơ sở hạ tầng. Đây là phần cứng thực hiện các thao tác vật lý.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Phần con người (Humanware - H): ").bold = True
    p.add_run("Kiến thức, kỹ năng, kinh nghiệm, thái độ, năng lực sáng tạo của người lao động tham gia vào quá trình vận hành và quản lý công nghệ.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Phần thông tin (Inforware - I): ").bold = True
    p.add_run("Các dữ liệu, bản vẽ kỹ thuật, tài liệu, quy trình, bí quyết (know-how), công thức, phần mềm. Nó giúp con người vận hành và kiểm soát máy móc.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Phần tổ chức (Orgaware - O): ").bold = True
    p.add_run("Cơ cấu tổ chức, quy định, mạng lưới phối hợp, thể chế quản lý để liên kết T, H, và I lại với nhau thành một hệ thống hoạt động hiệu quả.")

    # Câu 4
    add_heading(doc, "Câu 4: Phân tích các thành phần của Marketing Mix (4P)", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Product (Sản phẩm): ").bold = True
    p.add_run("Là những hàng hóa hoặc dịch vụ mà doanh nghiệp cung cấp cho thị trường mục tiêu nhằm thỏa mãn nhu cầu của khách hàng. Phân tích bao gồm: chất lượng, thiết kế, tính năng, nhãn hiệu, bao bì, kích cỡ, dịch vụ đi kèm, bảo hành.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Price (Giá cả): ").bold = True
    p.add_run("Là số tiền khách hàng phải trả để sở hữu hoặc sử dụng sản phẩm. Đây là chữ P duy nhất tạo ra doanh thu. Phân tích bao gồm: chiến lược định giá, giá niêm yết, chiết khấu, thời hạn thanh toán, điều kiện trả chậm.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Place (Phân phối): ").bold = True
    p.add_run("Là các hoạt động làm cho sản phẩm có mặt trên thị trường và tiếp cận được với người tiêu dùng mục tiêu. Phân tích bao gồm: kênh phân phối (trực tiếp, gián tiếp), độ bao phủ, vị trí điểm bán, vận chuyển, quản lý kho bãi, logistics.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Promotion (Xúc tiến thương mại/Truyền thông): ").bold = True
    p.add_run("Là các hoạt động thông tin, thuyết phục và nhắc nhở khách hàng về sản phẩm và doanh nghiệp. Phân tích bao gồm: quảng cáo, khuyến mãi, quan hệ công chúng (PR), bán hàng cá nhân, marketing trực tiếp.")

    doc.add_page_break()

    add_heading(doc, "PHẦN 2: BÀI TẬP TÍNH TOÁN", level=1)

    # Câu 5
    add_heading(doc, "Câu 5: Bài toán Tồn kho (EOQ cơ bản)", level=2)
    add_paragraph(doc, "Tóm tắt dữ kiện:", bold=True)
    add_paragraph(doc, "Nhu cầu hằng năm (D) = 24.000 vòng bi\nChi phí đặt hàng (S) = 600.000 đồng/lần\nChi phí tồn trữ (H) = 12.000 đồng/đơn vị/năm\nSố ngày làm việc = 300 ngày\nThời gian giao hàng (L) = 5 ngày\nDự trữ an toàn (SS) = 200 vòng bi")
    
    add_paragraph(doc, "Giải:", bold=True)
    add_paragraph(doc, "1. Lượng đặt hàng kinh tế (EOQ):")
    add_paragraph(doc, "EOQ = √(2.D.S / H) = √(2 × 24.000 × 600.000 / 12.000) = √2.400.000 = 1.549,19 ≈ 1.549 (vòng bi)")
    
    add_paragraph(doc, "2. Số lần đặt hàng trong năm (N):")
    add_paragraph(doc, "N = D / EOQ = 24.000 / 1.549 = 15,49 (lần/năm)")
    
    add_paragraph(doc, "3. Thời gian giữa hai lần đặt hàng (T):")
    add_paragraph(doc, "T = Số ngày làm việc / N = 300 / 15,49 = 19,36 (ngày)")
    
    add_paragraph(doc, "4. Tổng chi phí tồn kho hằng năm (TC):")
    add_paragraph(doc, "- Chi phí đặt hàng hằng năm = (D / EOQ) × S = 15,49 × 600.000 = 9.294.000 (đồng)")
    add_paragraph(doc, "- Chi phí tồn trữ hằng năm = (EOQ / 2) × H = (1.549 / 2) × 12.000 = 9.294.000 (đồng)")
    add_paragraph(doc, "=> Tổng chi phí (TC) = 9.294.000 + 9.294.000 = 18.588.000 (đồng)")
    
    add_paragraph(doc, "5. Điểm tái đặt hàng (ROP):")
    add_paragraph(doc, "Nhu cầu sử dụng mỗi ngày (d) = D / Số ngày làm việc = 24.000 / 300 = 80 (vòng bi/ngày)")
    add_paragraph(doc, "ROP = d × L + SS = 80 × 5 + 200 = 400 + 200 = 600 (vòng bi)")
    
    add_paragraph(doc, "6. Nếu thời gian giao hàng tăng lên 8 ngày:")
    add_paragraph(doc, "ROP mới = d × L_mới + SS = 80 × 8 + 200 = 640 + 200 = 840 (vòng bi).")
    add_paragraph(doc, "=> Điểm tái đặt hàng sẽ tăng thêm 240 vòng bi.")

    # Câu 6
    add_heading(doc, "Câu 6: Bài toán Tồn kho (Khấu trừ theo số lượng)", level=2)
    add_paragraph(doc, "Tóm tắt dữ kiện:", bold=True)
    add_paragraph(doc, "Nhu cầu (D) = 18.000 m\nChi phí đặt hàng (S) = 2.400.000 đồng/lần\nTỷ lệ chi phí tồn trữ (I) = 15% = 0,15\nSố ngày làm việc = 360 ngày\nThời gian giao hàng (L) = 6 ngày\nCác mức giá: P1 = 95.000 (Q < 4.000); P2 = 92.000 (4.000 ≤ Q < 8.000); P3 = 89.000 (Q ≥ 8.000)")
    
    add_paragraph(doc, "Giải:", bold=True)
    add_paragraph(doc, "Bước 1: Tính EOQ ở từng mức giá.")
    add_paragraph(doc, "Công thức: EOQi = √(2.D.S / (I.Pi))")
    add_paragraph(doc, "- EOQ3 (P3 = 89.000): √(2 × 18.000 × 2.400.000 / (0,15 × 89.000)) = √6.471.910 = 2.544 (m)")
    add_paragraph(doc, "- EOQ2 (P2 = 92.000): √(2 × 18.000 × 2.400.000 / (0,15 × 92.000)) = √6.260.869 = 2.502 (m)")
    add_paragraph(doc, "- EOQ1 (P1 = 95.000): √(2 × 18.000 × 2.400.000 / (0,15 × 95.000)) = √6.063.157 = 2.462 (m)")
    
    add_paragraph(doc, "Bước 2: Xác định mức đặt hàng hợp lý ở từng khoảng (Điều chỉnh EOQ).")
    add_paragraph(doc, "- Xét khoảng giá P3 (Q ≥ 8000): EOQ3 = 2.544 < 8000 (Không hợp lệ) => Điều chỉnh Q3* = 8.000")
    add_paragraph(doc, "- Xét khoảng giá P2 (4000 ≤ Q < 8000): EOQ2 = 2.502 < 4000 (Không hợp lệ) => Điều chỉnh Q2* = 4.000")
    add_paragraph(doc, "- Xét khoảng giá P1 (Q < 4000): EOQ1 = 2.462 < 4000 (Hợp lệ) => Q1* = 2.462")
    
    add_paragraph(doc, "Bước 3: Tính tổng chi phí (TC) cho các phương án hợp lý. (TC = D.P + (D/Q).S + (Q/2).I.P)")
    add_paragraph(doc, "- Tại Q3* = 8.000 (P=89.000):")
    add_paragraph(doc, "  TC3 = (18.000 × 89.000) + (18.000 / 8.000 × 2.400.000) + (8.000 / 2 × 0,15 × 89.000)")
    add_paragraph(doc, "      = 1.602.000.000 + 5.400.000 + 53.400.000 = 1.660.800.000 (đồng)")
    
    add_paragraph(doc, "- Tại Q2* = 4.000 (P=92.000):")
    add_paragraph(doc, "  TC2 = (18.000 × 92.000) + (18.000 / 4.000 × 2.400.000) + (4.000 / 2 × 0,15 × 92.000)")
    add_paragraph(doc, "      = 1.656.000.000 + 10.800.000 + 27.600.000 = 1.694.400.000 (đồng)")
    
    add_paragraph(doc, "- Tại Q1* = 2.462 (P=95.000):")
    add_paragraph(doc, "  TC1 = (18.000 × 95.000) + (18.000 / 2.462 × 2.400.000) + (2.462 / 2 × 0,15 × 95.000)")
    add_paragraph(doc, "      = 1.710.000.000 + 17.546.710 + 17.541.750 = 1.745.088.460 (đồng)")
    
    add_paragraph(doc, "Bước 4: Chọn quy mô đặt hàng tối ưu.")
    add_paragraph(doc, "So sánh TC, ta thấy TC3 < TC2 < TC1. Do đó, mức đặt hàng tối ưu là Q = 8.000 mét vải.")
    
    add_paragraph(doc, "Bước 5: Tính điểm tái đặt hàng (ROP) theo phương án được chọn.")
    add_paragraph(doc, "Nhu cầu mỗi ngày (d) = D / Số ngày làm việc = 18.000 / 360 = 50 (m/ngày)")
    add_paragraph(doc, "ROP = d × L = 50 × 6 = 300 (m)")
    
    add_paragraph(doc, "Bước 6: Nhận xét.")
    add_paragraph(doc, "=> Nên ưu tiên ĐẶT NHIỀU ĐỂ HƯỞNG CHIẾT KHẤU. (Vì tiền tiết kiệm được từ giá mua lớn hơn rất nhiều so với tiền lưu kho tăng thêm).", bold=True)

    doc.add_page_break()

    # Câu 7
    add_heading(doc, "Câu 7: Quản lý dự án - Sơ đồ mạng AON (CPM)", level=2)
    add_paragraph(doc, "Giải:", bold=True)
    add_paragraph(doc, "1. Sơ đồ mạng AON, ES, EF, LS, LF, Dự trữ (Slack):")
    
    # ASCII Diagram for Q7
    diagram_q7 = """
          +--- B(4) ---+--- D(6) ---+
          |            |            |
Start --- A(3)         +--- E(3) ---+--- G(5) ---+
          |                                      |
          +--- C(5) ------- F(4) ------- H(2) ---+--- I(4) --- End
"""
    p_diag = doc.add_paragraph()
    run_diag = p_diag.add_run(diagram_q7)
    run_diag.font.name = 'Courier New'
    run_diag.font.size = Pt(10)
    
    add_paragraph(doc, "Tính toán thông số mạng từng công việc:")
    add_paragraph(doc, "- Công việc A (t=3): ES = 0, EF = 0+3=3; LS = 0, LF = 3; Dự trữ S = 0 (Găng).")
    add_paragraph(doc, "- Công việc B (t=4): ES = 3, EF = 3+4=7; LS = 3, LF = 7; Dự trữ S = 0 (Găng).")
    add_paragraph(doc, "- Công việc C (t=5): ES = 3, EF = 3+5=8; LS = 8, LF = 13; Dự trữ S = 5.")
    add_paragraph(doc, "- Công việc D (t=6): ES = 7, EF = 7+6=13; LS = 7, LF = 13; Dự trữ S = 0 (Găng).")
    add_paragraph(doc, "- Công việc E (t=3): ES = 7, EF = 7+3=10; LS = 10, LF = 13; Dự trữ S = 3.")
    add_paragraph(doc, "- Công việc F (t=4): ES = 8, EF = 8+4=12; LS = 13, LF = 17; Dự trữ S = 5.")
    add_paragraph(doc, "- Công việc G (t=5): ES = 13, EF = 13+5=18; LS = 13, LF = 18; Dự trữ S = 0 (Găng).")
    add_paragraph(doc, "- Công việc H (t=2): ES = 12, EF = 12+2=14; LS = 17, LF = 19; Dự trữ S = 5.")
    add_paragraph(doc, "- Công việc I (t=4): ES = 18, EF = 18+4=22; LS = 18, LF = 22; Dự trữ S = 0 (Găng).")

    add_paragraph(doc, "\n2. Vẽ biểu đồ Gantt theo phương án triển khai sớm (ES):")
    gantt_q7 = """
(Mỗi dấu = tương ứng với 1 ngày thực hiện)
Ngày: 01 02 03 04 05 06 07 08 09 10 11 12 13 14 15 16 17 18 19 20 21 22
A   : ======\\
B   :       ========\\
C   :       ==========\\
D   :               ============\\
E   :               ======\\
F   :                 ========\\
G   :                           ==========\\
H   :                         ====\\
I   :                                     ========\\
"""
    p_gantt = doc.add_paragraph()
    run_gantt = p_gantt.add_run(gantt_q7)
    run_gantt.font.name = 'Courier New'
    run_gantt.font.size = Pt(10)
    
    add_paragraph(doc, "\n3. Đường găng:")
    add_paragraph(doc, "Các công việc có dự trữ (S) = 0 là: A, B, D, G, I")
    add_paragraph(doc, "=> Đường găng là: A -> B -> D -> G -> I")
    
    add_paragraph(doc, "3. Thời gian hoàn thành toàn dự án:")
    add_paragraph(doc, "T = 22 (ngày)")
    
    add_paragraph(doc, "4. Nếu công việc F bị chậm thêm 3 ngày thì dự án có chậm không? Giải thích.")
    add_paragraph(doc, "Công việc F không nằm trên đường găng và có thời gian dự trữ là S = 5 ngày. Nếu F bị chậm thêm 3 ngày (3 < 5) thì dự án sẽ KHÔNG bị chậm trễ vì độ chậm trễ vẫn nằm trong giới hạn thời gian dự trữ cho phép.")

    # Câu 8
    add_heading(doc, "Câu 8: Quản lý dự án - Sơ đồ mạng PERT", level=2)
    add_paragraph(doc, "Công thức: Thời gian kỳ vọng (t) = (a + 4m + b) / 6. Phương sai (V) = ((b - a) / 6)²")
    
    table2 = doc.add_table(rows=1, cols=7)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'CV'
    hdr_cells2[1].text = 'a'
    hdr_cells2[2].text = 'm'
    hdr_cells2[3].text = 'b'
    hdr_cells2[4].text = 't (Kỳ vọng)'
    hdr_cells2[5].text = 'Phương sai (V)'
    hdr_cells2[6].text = 'Đường trước'
    
    data2 = [
        ('A', 2, 4, 6, 4, 0.44, '-'),
        ('B', 3, 5, 7, 5, 0.44, 'A'),
        ('C', 2, 3, 4, 3, 0.11, 'A'),
        ('D', 4, 6, 8, 6, 0.44, 'B'),
        ('E', 1, 2, 3, 2, 0.11, 'C'),
        ('F', 3, 5, 7, 5, 0.44, 'D, E'),
        ('G', 2, 3, 4, 3, 0.11, 'F')
    ]
    
    for item in data2:
        row_cells = table2.add_row().cells
        for i in range(7):
            row_cells[i].text = str(item[i])
            
    add_paragraph(doc, "\n1 & 2. Đã tính ở bảng trên.")
    
    add_paragraph(doc, "3 & 4. Xác định đường găng và thời gian hoàn thành dự án.")
    add_paragraph(doc, "Sơ đồ mạng AON:")
    
    diagram_q8 = """
          +--- B(5) --- D(6) ---+
          |                     |
Start --- A(4)                  +--- F(5) --- G(3) --- End
          |                     |
          +--- C(3) --- E(2) ---+
"""
    p_diag_8 = doc.add_paragraph()
    run_diag_8 = p_diag_8.add_run(diagram_q8)
    run_diag_8.font.name = 'Courier New'
    run_diag_8.font.size = Pt(10)
    
    add_paragraph(doc, "Các đường có thể đi: A -> B -> D -> F -> G (đường 1) và A -> C -> E -> F -> G (đường 2)")
    add_paragraph(doc, "Tổng thời gian đường 1: tA + tB + tD + tF + tG = 4 + 5 + 6 + 5 + 3 = 23 (ngày)")
    add_paragraph(doc, "Tổng thời gian đường 2: tA + tC + tE + tF + tG = 4 + 3 + 2 + 5 + 3 = 17 (ngày)")
    add_paragraph(doc, "=> Đường găng là đường dài nhất: A -> B -> D -> F -> G")
    add_paragraph(doc, "Thời gian kỳ vọng hoàn thành dự án (T_E) = 23 (ngày).")
    
    add_paragraph(doc, "5. Tính xác suất hoàn thành dự án:")
    add_paragraph(doc, "Tổng phương sai của dự án (trên đường găng):")
    add_paragraph(doc, "V_p = V_A + V_B + V_D + V_F + V_G = 0.44 + 0.44 + 0.44 + 0.44 + 0.11 = 1.87")
    add_paragraph(doc, "Độ lệch chuẩn (σ) = √1.87 ≈ 1.367")
    
    add_paragraph(doc, "a) Hoàn thành trong 20 ngày:")
    add_paragraph(doc, "Z = (T - T_E) / σ = (20 - 23) / 1.367 = -2.19")
    add_paragraph(doc, "Tra bảng phân phối chuẩn: P(Z ≤ -2.19) ≈ 1.43% (Xác suất rất thấp)")
    
    add_paragraph(doc, "b) Hoàn thành trong 22 ngày:")
    add_paragraph(doc, "Z = (22 - 23) / 1.367 = -0.73")
    add_paragraph(doc, "Tra bảng: P(Z ≤ -0.73) ≈ 23.27%")
    
    add_paragraph(doc, "6. Muốn đạt xác suất 95%:")
    add_paragraph(doc, "Tra bảng phân phối chuẩn ngược: P(Z) = 0.95 => Z ≈ 1.645 (hoặc 1.65)")
    add_paragraph(doc, "Z = (T - T_E) / σ => 1.645 = (T - 23) / 1.367")
    add_paragraph(doc, "=> T = 23 + (1.645 × 1.367) = 23 + 2.25 = 25.25 (ngày)")
    add_paragraph(doc, "Vậy doanh nghiệp cần dự trù tối thiểu khoảng 25,25 ngày (khoảng 26 ngày).")

    # Save document
    out_path = os.path.join(os.getcwd(), 'docs', 'QTKDCKS', 'Giai_chi_tiet_De_cuong_QTKDCKS.docx')
    doc.save(out_path)
    print(f"File saved successfully to {out_path}")

if __name__ == '__main__':
    main()