from docx import Document
from docx.shared import Pt

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('HƯỚNG DẪN TRÌNH BÀY GIẤY THI - AN TOÀN VÀ BẢO MẬT THÔNG TIN', 0)
    doc.add_paragraph('(Chép y nguyên theo từng bước dưới đây vào giấy thi để đạt điểm tối đa)')
    
    # CÂU 1
    doc.add_heading('CÂU 1: GIẢI MÃ AFFINE', level=1)
    doc.add_paragraph('Đề bài: Giải mã "rmpjvpwtwtmhgmpwtgfwgkzn" với a=19, b=23.')
    doc.add_heading('Giải:', level=2)
    doc.add_paragraph('Bước 1: Xác định công thức giải mã:\n- Hệ mã Affine trên Z_26.\n- Công thức giải mã: P = a^(-1) * (C - b) mod 26')
    doc.add_paragraph('Bước 2: Tìm phần tử nghịch đảo của a=19:\n- Ta cần tìm a^(-1) sao cho 19 * a^(-1) ≡ 1 mod 26\n- 19 * 11 = 209 = 8 * 26 + 1\n- Suy ra a^(-1) = 11.')
    doc.add_paragraph('Bước 3: Lập công thức giải mã cụ thể:\n- P = 11 * (C - 23) mod 26\n- Do -23 ≡ 3 mod 26 nên P = 11 * (C + 3) mod 26')
    doc.add_paragraph('Bước 4: Giải mã bản mã:\nC=17 (r) -> P = 11*(17+3) mod 26 = 12 (M)\nC=12 (m) -> P = 11*(12+3) mod 26 = 9 (J)\nC=15 (p) -> P = 11*(15+3) mod 26 = 16 (Q)\n...\nBản rõ thu được: MJQCEQPIPIJGVJQPIVKVWNU')
    
    # CÂU 2
    doc.add_heading('CÂU 2: MÃ HILL (MA TRẬN 2x2)', level=1)
    doc.add_paragraph('Đề bài: Mã hóa "doikhongnhulamo" với khóa K = [[7, 3], [8, 7]]. Giải mã lại.')
    doc.add_heading('A. MÃ HÓA', level=2)
    doc.add_paragraph('Bước 1: Chuẩn bị bản rõ\n- Bản rõ: do ik ho ng nh ul am ox (thêm "x" ở cuối để chẵn)\n- Đổi sang số: (3,14), (8,10), (7,14), (13,6), (13,7), (20,11), (0,12), (14,23)')
    doc.add_paragraph('Bước 2: Mã hóa C = P x K mod 26\n- Block 1 (do): [3, 14] x K = [3*7 + 14*8, 3*3 + 14*7] = [133, 107] ≡ [3, 3] mod 26 => DD\n- Block 2 (ik): [8, 10] x K = [136, 94] ≡ [6, 16] mod 26 => GQ\n- Block 3 (ho): [7, 14] x K = [161, 119] ≡ [5, 15] mod 26 => FP\n- Block 4 (ng): [13, 6] x K = [139, 81] ≡ [9, 3] mod 26 => JD\n- Block 5 (nh): [13, 7] x K = [147, 88] ≡ [17, 10] mod 26 => RK\n- Block 6 (ul): [20, 11] x K = [228, 137] ≡ [20, 7] mod 26 => UH\n- Block 7 (am): [0, 12] x K = [96, 84] ≡ [18, 6] mod 26 => SG\n- Block 8 (ox): [14, 23] x K = [282, 203] ≡ [22, 21] mod 26 => WV\nBản mã thu được: DD GQ FP JD RK UH SG WV')

    doc.add_heading('B. GIẢI MÃ', level=2)
    doc.add_paragraph('Bước 1: Tính định thức det(K)\n- det(K) = 7*7 - 3*8 = 49 - 24 = 25 ≡ 25 mod 26')
    doc.add_paragraph('Bước 2: Tìm det(K)^(-1)\n- Vì 25 ≡ -1 mod 26, nên (-1) * (-1) = 1 => det(K)^(-1) = 25 mod 26')
    doc.add_paragraph('Bước 3: Tính ma trận nghịch đảo K^(-1)\n- Đổi vị trí chéo chính, đổi dấu chéo phụ: Adj(K) = [[7, -3], [-8, 7]] ≡ [[7, 23], [18, 7]] mod 26\n- K^(-1) = 25 * [[7, 23], [18, 7]] = [[175, 575], [450, 175]] ≡ [[19, 3], [8, 19]] mod 26')
    doc.add_paragraph('Bước 4: Giải mã P = C x K^(-1) mod 26\n- Block 1 (DD): [3, 3] x [[19, 3], [8, 19]] = [3*19 + 3*8, 3*3 + 3*19] = [81, 66] ≡ [3, 14] mod 26 => do\n(Thực hiện tương tự cho các block tiếp theo ta thu lại được bản rõ ban đầu: doikhongnhulamo)')

    doc.add_heading('CÂU 4: MÃ HÓA RSA', level=1)
    doc.add_paragraph('Bước 1: Khởi tạo thông số RSA\n- n = p * q = 11 * 13 = 143\n- φ(n) = (p-1)*(q-1) = 10 * 12 = 120')
    doc.add_paragraph('Bước 2: Chọn e và tính khóa bí mật d\n- Chọn e = 7 (vì gcd(7, 120) = 1). Khóa công khai PU = (7, 143).\n- Tính d ≡ 7^(-1) mod 120. Do 7 * 103 = 721 = 120 * 6 + 1 => d = 103. Khóa bí mật PR = (103, 143).')
    doc.add_paragraph('Bước 3: Mã hóa C = M^e mod 143\n- Chữ c = 2 -> C1 = 2^7 mod 143 = 128\n- Chữ h = 7 -> C2 = 7^7 mod 143 = 6\n- Chữ u = 20 -> C3 = 20^7 mod 143 = 125\n...')

    doc.add_heading('CÂU 6: MÃ HÓA PLAYFAIR', level=1)
    doc.add_paragraph('Bước 1: Lập bảng Playfair 5x5 từ khóa "bainaylamroi"\nB A I N Y\nL M R O C\nD E F G H\nK P Q S T\nU V W X Z')
    doc.add_paragraph('Bước 2: Mã hóa tên LE VAN AN\n- Chia cặp: LE - VA - NA - NX\n- LE: Góc hình chữ nhật -> MD\n- VA: Cùng cột -> AV\n- NA: Cùng hàng -> YI\n- NX: Cùng cột -> OS\nBản mã: MD AV YI OS')

    doc.add_page_break()
    
    doc.add_heading('LÝ THUYẾT (HỌC THUỘC CHÉP VÀO LÀ ĂN ĐIỂM)', level=1)
    doc.add_heading('Câu 5 & 8: IDS và IPS', level=2)
    doc.add_paragraph('- IDS (Phát hiện xâm nhập): Thụ động. Chỉ quan sát mạng và gửi báo động (Alert), không trực tiếp chặn.\n- IPS (Ngăn ngừa xâm nhập): Chủ động. Nằm trực tiếp trên luồng mạng. Nếu phát hiện thì cắt kết nối (Drop/Block) gói tin ngay.\n- Cơ chế phát hiện: Dựa trên chữ ký (Signature) hoặc dựa trên sự bất thường (Anomaly).')
    
    doc.add_heading('Câu 9: Vai trò SSL/TLS', level=2)
    doc.add_paragraph('1. Mã hóa (Encryption): Tránh bị nghe trộm, biến dữ liệu thành bản mã khó đọc.\n2. Xác thực (Authentication): Đảm bảo kết nối đúng máy chủ thông qua chứng chỉ số (Certificate).\n3. Toàn vẹn (Integrity): Đảm bảo dữ liệu không bị kẻ gian sửa đổi giữa đường.')

    doc.add_heading('Câu 10: Chữ ký số', level=2)
    doc.add_paragraph('- Người gửi (Ký): Dùng hàm Băm (Hash) tóm tắt tài liệu -> Dùng Khóa bí mật (Private Key) để mã hóa bản băm -> Thành Chữ ký số đính kèm.\n- Người nhận (Kiểm tra): Dùng Khóa công khai (Public Key) giải mã chữ ký số -> So sánh với bản băm mới để xác thực.\n- Tác dụng: Chống giả mạo, chống chối bỏ và đảm bảo tính toàn vẹn.')

    doc.save('docs/BMTT/Giai_Viet_Tay_Sieu_De_Nhin.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()
