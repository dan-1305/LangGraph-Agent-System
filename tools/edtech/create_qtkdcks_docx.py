from docx import Document
from docx.shared import Pt

def create_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    doc.add_heading('HƯỚNG DẪN GIẢI CHI TIẾT - ĐỀ CƯƠNG QTKDCKS', 0)
    doc.add_paragraph('Tài liệu ôn thi bao gồm 5 câu lý thuyết và 14 câu bài tập thực hành được trình bày chi tiết.')

    # PHẦN 1: LÝ THUYẾT
    doc.add_heading('PHẦN 1: CÂU HỎI LÝ THUYẾT', level=1)
    
    doc.add_heading('Câu 1: Các bước trong quy trình ra quyết định quản lý (Tham khảo: Chương 2 - Ra quyết định, trang 14-16)', level=2)
    doc.add_paragraph('1. Nhận diện vấn đề: Xác định tình huống hoặc thách thức cần giải quyết.')
    doc.add_paragraph('2. Xác định tiêu chuẩn: Đưa ra các mốc/chỉ tiêu để đánh giá (ví dụ: chi phí, thời gian, chất lượng).')
    doc.add_paragraph('3. Đánh giá trọng số: Phân bổ mức độ quan trọng cho từng tiêu chuẩn.')
    doc.add_paragraph('4. Phát triển các phương án: Liệt kê tất cả các giải pháp khả thi.')
    doc.add_paragraph('5. Lựa chọn phương án: Đánh giá và chọn ra phương án tối ưu nhất dựa trên các tiêu chuẩn.')
    doc.add_paragraph('6. Thực thi và đánh giá: Triển khai phương án đã chọn và đo lường tính hiệu quả.')

    doc.add_heading('Câu 2: Phân biệt quyết định trong điều kiện chắc chắn, rủi ro và không chắc chắn (Tham khảo: Chương 2 - Ra quyết định, trang 23)', level=2)
    doc.add_paragraph('- Chắc chắn: Người ra quyết định biết chính xác 100% kết quả của từng phương án sẽ xảy ra (Môi trường hoàn toàn tĩnh).')
    doc.add_paragraph('- Rủi ro: Không biết chắc kết quả, nhưng biết được "xác suất" (tỷ lệ phần trăm) xảy ra của từng trường hợp (Môi trường có biến động nhưng đo lường được).')
    doc.add_paragraph('- Không chắc chắn: Hoàn toàn mù mờ, không có thông tin và cũng KHÔNG BIẾT xác suất xảy ra của các trường hợp (Môi trường hoàn toàn vô định).')

    doc.add_heading('Câu 3: Cấu trúc bảng cân đối kế toán và mối quan hệ Tài sản – Nợ – Vốn (Tham khảo: Chương 4 - Kế toán tài chính, trang 8-12)', level=2)
    doc.add_paragraph('- Tài sản: Toàn bộ nguồn lực mang lại lợi ích kinh tế mà doanh nghiệp đang sở hữu (Tiền mặt, Hàng tồn kho, Máy móc, Bất động sản...).')
    doc.add_paragraph('- Nguồn vốn: Nơi cung cấp nguồn lực để hình thành nên Tài sản. Bao gồm:')
    doc.add_paragraph('  + Nợ phải trả: Vay mượn từ bên ngoài (Ngân hàng, nhà cung cấp...).')
    doc.add_paragraph('  + Vốn chủ sở hữu: Tiền tự có của chủ doanh nghiệp hoặc cổ đông.')
    doc.add_paragraph('- Mối quan hệ cốt lõi: TỔNG TÀI SẢN = TỔNG NỢ + VỐN CHỦ SỞ HỮU.')

    doc.add_heading('Câu 4: Cấu trúc báo cáo kết quả kinh doanh và ý nghĩa (Tham khảo: Chương 4 - Kế toán tài chính, trang 15-18)', level=2)
    doc.add_paragraph('- Cấu trúc: Đi từ Doanh thu thuần -> Trừ Giá vốn hàng bán -> Lãi gộp -> Trừ Chi phí hoạt động -> Lợi nhuận trước thuế (LNTT) -> Trừ Thuế -> Lợi nhuận sau thuế (LNST).')
    doc.add_paragraph('- Ý nghĩa: Cho biết bức tranh "Lời - Lỗ" của doanh nghiệp trong một khoảng thời gian nhất định (thường là quý hoặc năm). Đánh giá trực tiếp hiệu quả hoạt động kinh doanh.')

    doc.add_heading('Câu 5: Ý nghĩa các tỷ số tài chính cơ bản (Tham khảo: Chương 4 - Kế toán tài chính, phần phân tích tỷ số, trang 25-30)', level=2)
    doc.add_paragraph('1. Tỷ số nợ (= Tổng nợ / Tổng tài sản): Cho biết bao nhiêu phần trăm tài sản được tài trợ bằng tiền đi vay. Tỷ lệ càng cao, rủi ro tài chính càng lớn.')
    doc.add_paragraph('2. Tỷ số thanh toán hiện hành (= TS lưu động / Nợ ngắn hạn): Đo lường khả năng trả các khoản nợ sắp đến hạn bằng tài sản ngắn hạn. (Tốt nhất > 1).')
    doc.add_paragraph('3. Tỷ số thanh toán nhanh (= (TS lưu động - Hàng tồn kho) / Nợ ngắn hạn): Giống tỷ số trên nhưng loại bỏ Hàng tồn kho vì HTK khó quy đổi ra tiền mặt ngay lập tức. Đánh giá tính thanh khoản khắt khe hơn.')
    doc.add_paragraph('4. Vòng quay tài sản (= Doanh thu / Tổng tài sản): Đo lường hiệu suất sử dụng tài sản để tạo ra doanh thu. Số vòng quay càng lớn, quản lý tài sản càng hiệu quả.')

    doc.add_page_break()

    # PHẦN 2: BÀI TẬP
    doc.add_heading('PHẦN 2: BÀI TẬP THỰC HÀNH', level=1)

    doc.add_heading('Bài 1: Ra quyết định rủi ro (EMV, EOL, EVPI) (Tham khảo: Chương 2 - Bài tập Ra quyết định, trang 28-35)', level=2)
    doc.add_paragraph('Đề bài: Có 3 phương án A (300; -100), B (200; -50), C (0; 0). Xác suất (0.6; 0.4).')
    doc.add_paragraph('1. Tính giá trị kỳ vọng (EMV) và chọn phương án:')
    doc.add_paragraph('  - EMV(A) = 300 * 0.6 + (-100) * 0.4 = 180 - 40 = 140')
    doc.add_paragraph('  - EMV(B) = 200 * 0.6 + (-50) * 0.4 = 120 - 20 = 100')
    doc.add_paragraph('  - EMV(C) = 0 * 0.6 + 0 * 0.4 = 0')
    doc.add_paragraph('=> Trả lời: Chọn phương án A vì có EMV cao nhất (140).')
    doc.add_paragraph('2. Tính tổn thất cơ hội kỳ vọng (EOL):')
    doc.add_paragraph('  - Ma trận tiếc nuối (EOL Table):')
    doc.add_paragraph('    + Thị trường Tốt (Max=300): A tiếc (300-300=0), B tiếc (300-200=100), C tiếc (300-0=300).')
    doc.add_paragraph('    + Thị trường Xấu (Max=0): A tiếc (0-(-100)=100), B tiếc (0-(-50)=50), C tiếc (0-0=0).')
    doc.add_paragraph('  - Tính EOL:')
    doc.add_paragraph('    + EOL(A) = 0 * 0.6 + 100 * 0.4 = 40')
    doc.add_paragraph('    + EOL(B) = 100 * 0.6 + 50 * 0.4 = 60 + 20 = 80')
    doc.add_paragraph('    + EOL(C) = 300 * 0.6 + 0 * 0.4 = 180')
    doc.add_paragraph('=> Trả lời: Chọn phương án A vì có EOL thấp nhất (Min EOL = 40).')
    doc.add_paragraph('3. Tính EVPI (Giá trị thông tin hoàn hảo):')
    doc.add_paragraph('  - Công thức: EVPI = Min EOL = 40 triệu đồng.')
    doc.add_paragraph('  - Kết luận: Giá mua thông tin (20 triệu) < EVPI (40 triệu). => Nên mua thông tin vì mang lại lợi ích kinh tế lớn hơn chi phí bỏ ra.')

    doc.add_heading('Bài 2: Ra quyết định không chắc chắn (Tham khảo: Chương 2 - Bài tập Ra quyết định, trang 24-27)', level=2)
    doc.add_paragraph('Đề bài: Trò 1 (40; -20), Trò 2 (60; -50).')
    doc.add_paragraph('a. Tiêu chuẩn Maximax (Lạc quan - Chọn cái lời nhất trong những cái lời nhất):')
    doc.add_paragraph('  - Lời Max của Trò 1: 40. Lời Max của Trò 2: 60. => Chọn Trò 2.')
    doc.add_paragraph('b. Tiêu chuẩn Maximin (Bi quan - Chọn cái ít lỗ nhất trong những cái lỗ nặng nhất):')
    doc.add_paragraph('  - Lỗ (Min) của Trò 1: -20. Lỗ (Min) của Trò 2: -50. Max(-20, -50) = -20. => Chọn Trò 1.')
    doc.add_paragraph('c. Tiêu chuẩn Hurwicz (α = 0.7):')
    doc.add_paragraph('  - Giá trị H(1) = 40 * 0.7 + (-20) * 0.3 = 28 - 6 = 22')
    doc.add_paragraph('  - Giá trị H(2) = 60 * 0.7 + (-50) * 0.3 = 42 - 15 = 27')
    doc.add_paragraph('=> Trả lời: Chọn Trò 2 vì có giá trị Hurwicz cao hơn (27 > 22).')

    doc.add_heading('Bài 3: Tính EMV với 3 trạng thái (Tham khảo: Chương 2 - Bài tập Ra quyết định, trang 28-35)', level=2)
    doc.add_paragraph('Đề bài: Xác suất (Xấu 0.2; TB 0.5; Tốt 0.3).')
    doc.add_paragraph('- EMV(A) = -100 * 0.2 + 50 * 0.5 + 150 * 0.3 = -20 + 25 + 45 = 50')
    doc.add_paragraph('- EMV(B) = -50 * 0.2 + 40 * 0.5 + 120 * 0.3 = -10 + 20 + 36 = 46')
    doc.add_paragraph('- EMV(C) = -20 * 0.2 + 20 * 0.5 + 80 * 0.3 = -4 + 10 + 24 = 30')
    doc.add_paragraph('=> Trả lời: Chọn phương án A vì mang lại lợi nhuận kỳ vọng cao nhất (EMV = 50).')

    doc.add_heading('Bài 4: Báo cáo kết quả kinh doanh (Tham khảo: Chương 4 - Kế toán tài chính, trang 15-18)', level=2)
    doc.add_paragraph('1. Lãi gộp = Doanh thu - Giá vốn = 10.000 - (60% * 10.000) = 4.000')
    doc.add_paragraph('2. Lợi nhuận trước thuế (LNTT) = Lãi gộp - Chi phí hoạt động - Lãi vay = 4.000 - 2.000 - 500 = 1.500')
    doc.add_paragraph('3. Lợi nhuận sau thuế (LNST) = LNTT - Thuế = 1.500 - (20% * 1.500) = 1.500 - 300 = 1.200')

    doc.add_heading('Bài 5: Tỷ số Nợ (Tham khảo: Chương 4 - Kế toán tài chính, trang 25)', level=2)
    doc.add_paragraph('1. Vốn chủ sở hữu = Tổng tài sản - Tổng nợ = 500.000 - 200.000 = 300.000')
    doc.add_paragraph('2. Tỷ số nợ = Tổng nợ / Tổng tài sản = 200.000 / 500.000 = 0.4 (40%)')

    doc.add_heading('Bài 6: Khả năng thanh toán (Tham khảo: Chương 4 - Kế toán tài chính, trang 26)', level=2)
    doc.add_paragraph('1. Tỷ số thanh toán hiện hành (Current ratio) = Tài sản lưu động / Nợ ngắn hạn = 300.000 / 150.000 = 2')
    doc.add_paragraph('2. Tỷ số thanh toán nhanh (Quick ratio) = (Tài sản lưu động - Hàng tồn kho) / Nợ ngắn hạn = (300.000 - 100.000) / 150.000 = 200.000 / 150.000 = 1.33')

    doc.add_heading('Bài 7 & 8: Chỉ số hoạt động (Tham khảo: Chương 4 - Kế toán tài chính, trang 27-28)', level=2)
    doc.add_paragraph('Bài 7: Vòng quay tài sản = Doanh thu / Tổng tài sản = 1.200.000 / 600.000 = 2 (vòng)')
    doc.add_paragraph('Bài 8: Kỳ thu tiền bình quân = Khoản phải thu / Doanh thu bình quân ngày = 120.000 / 4.000 = 30 (ngày)')

    doc.add_heading('Bài 9: Lợi nhuận giữ lại (Tham khảo: Chương 4 - Kế toán tài chính, Bảng cân đối kết nối KQKD)', level=2)
    doc.add_paragraph('Lợi nhuận giữ lại (LNGL) cuối kỳ = LNGL đầu kỳ + Lợi nhuận sau thuế (LNST) - Cổ tức chi trả')
    doc.add_paragraph('=> LNGL cuối kỳ = 50.000 + 80.000 - 20.000 = 110.000')

    doc.add_heading('Bài 10 & 11: Mô hình quản lý hàng tồn kho (EOQ) (Tham khảo: Chương 5 - Quản trị hàng tồn kho, mô hình EOQ)', level=2)
    doc.add_paragraph('Bài 10:')
    doc.add_paragraph('- Điểm đặt hàng tối ưu (EOQ) = Căn bậc hai của [(2 * D * S) / H] = Căn bậc hai [(2 * 10.000 * 50) / 5] = Căn bậc hai(200.000) = 447.21 (đơn vị)')
    doc.add_paragraph('- Số lần đặt hàng (N) = D / EOQ = 10.000 / 447.21 = 22.36 (lần)')
    doc.add_paragraph('- Chu kỳ đặt hàng (T) = 365 / N = 365 / 22.36 = 16.32 (ngày)')
    doc.add_paragraph('- Tổng chi phí tồn kho (TC) = Chi phí đặt hàng + Chi phí lưu kho = (D/EOQ)*S + (EOQ/2)*H = (22.36 * 50) + (223.6 * 5) = 1118 + 1118 = 2236 (đơn vị tiền tệ)')
    doc.add_paragraph('Bài 11:')
    doc.add_paragraph('- EOQ = Căn bậc hai của [(2 * 12.000 * 100) / 2] = Căn bậc hai(1.200.000) = 1095.45 (đơn vị)')
    doc.add_paragraph('- Nhu cầu bình quân 1 ngày làm việc (d) = Nhu cầu năm / Số ngày làm việc = 12.000 / 300 = 40 (đơn vị/ngày)')
    doc.add_paragraph('- Điểm đặt hàng lại (ROP) = d * Thời gian giao hàng (Lead time) = 40 * 5 = 200 (đơn vị).')

    doc.add_heading('Bài 12: Điểm đặt hàng lại có Dự trữ an toàn (Tham khảo: Chương 5 - Quản trị hàng tồn kho, mô hình ROP)', level=2)
    doc.add_paragraph('- ROP = (Nhu cầu 1 ngày * Lead time) + Dự trữ an toàn')
    doc.add_paragraph('- ROP = (40 * 4) + 50 = 160 + 50 = 210 (đơn vị). (Khi kho tụt xuống còn 210 món thì phải gọi điện đặt mua đợt mới).')

    doc.add_heading('Bài 13: Quản trị dự án (Đường găng CPM) (Tham khảo: Chương 7 - Quản trị dự án, phương pháp mạng PERT/CPM)', level=2)
    doc.add_paragraph('1. Phân tích các nhánh (đường đi) của dự án từ lúc bắt đầu đến kết thúc:')
    doc.add_paragraph('- Nhánh 1: A -> B -> D. Tổng thời gian: 3 + 4 + 6 = 13 (ngày)')
    doc.add_paragraph('- Nhánh 2: A -> C -> E. Tổng thời gian: 3 + 5 + 2 = 10 (ngày)')
    doc.add_paragraph('2. Thời gian hoàn thành dự án: Chính là độ dài của nhánh dài nhất = 13 ngày.')
    doc.add_paragraph('3. Đường găng (Đường tới hạn): Là nhánh có thời gian dài nhất = A -> B -> D.')
    doc.add_paragraph('4. Thời gian dự trữ (Slack/Float):')
    doc.add_paragraph('- Các công việc NẰM TRÊN đường găng (A, B, D) không có thời gian dự trữ. (Slack = 0). Nếu trễ 1 ngày dự án sẽ trễ theo.')
    doc.add_paragraph('- Các công việc KHÔNG nằm trên đường găng (C, E) có thời gian dự trữ = 13 - 10 = 3 ngày. (C và E có thể rề rà chậm 3 ngày mà không làm trễ dự án).')

    doc.add_heading('Bài 14: Quản trị dự án (Phương pháp PERT) (Tham khảo: Chương 7 - Quản trị dự án, phân tích phương sai)', level=2)
    doc.add_paragraph('Với a = Lạc quan (2), m = Bình thường (4), b = Bi quan (8).')
    doc.add_paragraph('1. Thời gian kỳ vọng (Te):')
    doc.add_paragraph('- Công thức: Te = (a + 4*m + b) / 6')
    doc.add_paragraph('- Te = (2 + 4*4 + 8) / 6 = (2 + 16 + 8) / 6 = 26 / 6 = 4.33')
    doc.add_paragraph('2. Phương sai (V):')
    doc.add_paragraph('- Công thức: V = ((b - a) / 6)^2')
    doc.add_paragraph('- V = ((8 - 2) / 6)^2 = (6 / 6)^2 = 1^2 = 1.')

    doc.save('docs/QTKDCKS/Giai_Chi_Tiet_QTKDCKS.docx')
    print("Done")

if __name__ == '__main__':
    create_docx()
